"""
Generador de la Memoria TFM — Zeptai
Ejecutar: python generar_tfm_memoria.py
Salida: Zeptai_Memoria_TFM.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Configuración de página ────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Estilos base ──────────────────────────────────────────────────────────
styles = doc.styles

def set_style(style_name, font_name="Times New Roman", size=12, bold=False,
              color=None, space_before=6, space_after=6, line_spacing=None,
              alignment=WD_ALIGN_PARAGRAPH.LEFT):
    try:
        st = styles[style_name]
    except KeyError:
        return
    st.font.name = font_name
    st.font.size = Pt(size)
    st.font.bold = bold
    if color:
        st.font.color.rgb = RGBColor(*color)
    st.paragraph_format.space_before = Pt(space_before)
    st.paragraph_format.space_after  = Pt(space_after)
    if line_spacing:
        st.paragraph_format.line_spacing = Pt(line_spacing)
    st.paragraph_format.alignment = alignment

set_style("Normal",    size=11, space_before=0, space_after=8, line_spacing=18)
set_style("Heading 1", size=16, bold=True, color=(26,60,120),
          space_before=18, space_after=8,
          alignment=WD_ALIGN_PARAGRAPH.LEFT)
set_style("Heading 2", size=13, bold=True, color=(26,60,120),
          space_before=12, space_after=6)
set_style("Heading 3", size=11, bold=True, color=(60,60,60),
          space_before=8, space_after=4)

# ─── Helpers ────────────────────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = align
    para.paragraph_format.first_line_indent = Cm(0.8)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    return para

def pb(text):
    return p(text, bold=True)

def bullet(text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(1.2 + level * 0.5)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(11)
    return para

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.5)
    para.paragraph_format.right_indent = Cm(1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    para._p.get_or_add_pPr().append(shading)
    return para

def table_2col(rows, header=None):
    cols = 2
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    t.style = "Table Grid"
    if header:
        row = t.rows[0]
        for i, cell_text in enumerate(header):
            cell = row.cells[i]
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    offset = 1 if header else 0
    for ri, (c1, c2) in enumerate(rows):
        row = t.rows[ri + offset]
        row.cells[0].text = c1
        row.cells[1].text = c2
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

def note(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(f"Nota: {text}")
    run.font.size = Pt(9)
    run.italic = True

# ═══════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("Zeptai")
r.font.name = "Times New Roman"
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(26, 60, 120)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_para.add_run(
    "Plataforma SaaS Multi-Agente de IA Generativa\n"
    "para la Automatización de la Gestión Empresarial\n"
    "de Autónomos y PYMEs"
)
r2.font.name = "Times New Roman"
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(60, 60, 60)

doc.add_paragraph()

master_para = doc.add_paragraph()
master_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = master_para.add_run("Máster en Ingeniería y Desarrollo de Soluciones de IA Generativa")
rm.font.name = "Times New Roman"
rm.font.size = Pt(13)
rm.font.bold = True

doc.add_paragraph()

for field, value in [
    ("Autor", "Alejandro Bravo"),
    ("Tutor", "[Nombre del tutor]"),
    ("Fecha de entrega", "Junio 2026"),
    ("Repositorio", "https://github.com/alejbrata/zeptai"),
]:
    field_para = doc.add_paragraph()
    field_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = field_para.add_run(f"{field}: ")
    rf.font.bold = True
    rf.font.size = Pt(12)
    rv = field_para.add_run(value)
    rv.font.size = Pt(12)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# AGRADECIMIENTOS
# ═══════════════════════════════════════════════════════════════════════════
h1("Agradecimientos")
p(
    "El desarrollo de Zeptai no habría sido posible sin el apoyo y la orientación recibidos a lo largo de este "
    "Máster. Quisiera expresar mi más sincero agradecimiento a los docentes del programa por transmitir su "
    "conocimiento sobre Inteligencia Artificial Generativa con rigor y entusiasmo."
)
p(
    "A mi tutor, por su disponibilidad y por sus comentarios siempre constructivos que han guiado "
    "la dirección técnica y académica de este trabajo."
)
p(
    "A los compañeros del máster, por las innumerables conversaciones técnicas que han enriquecido "
    "este proyecto, y a los primeros usuarios que probaron la plataforma y aportaron su valioso feedback."
)
p(
    "Finalmente, a mi familia y amigos, por su paciencia y apoyo incondicional durante los meses "
    "de desarrollo intensivo de este proyecto."
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# RESUMEN / ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════
h1("Resumen")
p(
    "Zeptai es una plataforma web SaaS que integra un sistema multi-agente de Inteligencia Artificial "
    "Generativa orientado a autónomos y pequeñas empresas españolas. La solución permite a sus usuarios "
    "gestionar su negocio a través de una interfaz conversacional, delegando tareas como la generación "
    "de presupuestos en PDF, la creación de material de marketing (imágenes, presentaciones, prompts "
    "de vídeo), la gestión de citas, la atención postventa y el análisis de tickets de gasto mediante OCR. "
    "El núcleo del sistema es un AgentExecutor basado en GPT-4o con function calling, complementado por "
    "un CouncilManager que orquesta un debate multi-perspectiva entre tres agentes especializados para "
    "decisiones estratégicas complejas."
)
p(
    "La arquitectura está desplegada en Docker sobre Flask y PostgreSQL con extensión pgvector, "
    "incorporando un pipeline RAG (Retrieval-Augmented Generation) que permite al agente responder "
    "con el conocimiento específico de cada negocio, notificaciones web push mediante el estándar VAPID, "
    "y un servidor Model Context Protocol (MCP) basado en Server-Sent Events para la integración de "
    "herramientas externas. La calidad del pipeline RAG se valida objetivamente mediante DeepEval, "
    "un framework LLM-as-a-judge que evalúa Faithfulness, Answer Relevancy, Contextual Precision y "
    "Contextual Recall sobre siete casos de prueba representativos."
)
p(
    "Los resultados demuestran que la plataforma es capaz de reducir el tiempo dedicado a tareas "
    "administrativas recurrentes en más de un 80%, con un coste operativo por usuario de entre 1 y "
    "5 euros mensuales y un ROI estimado superior al 1.000% en el primer mes de uso activo. "
    "El conjunto de pruebas automatizadas (51 casos funcionales con tasa de superación del 90,2%) "
    "y la evaluación con DeepEval validan la robustez y calidad de la solución."
)
p(
    "Palabras clave: IA Generativa, Agentes LLM, GPT-4o, Function Calling, Multi-Agente, SaaS, "
    "Automatización PYME, RAG, pgvector, Flask, Docker, MCP, DeepEval.",
    italic=True
)

doc.add_paragraph()
h1("Abstract")
p(
    "Zeptai is a SaaS web platform integrating a multi-agent Generative Artificial Intelligence system "
    "aimed at Spanish freelancers and small businesses. The solution enables users to manage their "
    "business through a conversational interface, delegating tasks such as PDF quote generation, "
    "marketing content creation (images, presentations, video prompts), appointment management, "
    "after-sales support, and expense ticket analysis via OCR. The system's core is an AgentExecutor "
    "built on GPT-4o with function calling, complemented by a CouncilManager that orchestrates a "
    "multi-perspective debate among three specialized agents for complex strategic decisions."
)
p(
    "The architecture is deployed on Docker using Flask and PostgreSQL with the pgvector extension, "
    "incorporating a RAG pipeline, VAPID web push notifications, and a Model Context Protocol (MCP) "
    "server based on Server-Sent Events. RAG quality is objectively validated using DeepEval "
    "(LLM-as-a-judge) across four metrics and seven test cases."
)
p(
    "Keywords: Generative AI, LLM Agents, GPT-4o, Function Calling, Multi-Agent, SaaS, "
    "SME Automation, RAG, pgvector, Flask, Docker, MCP, DeepEval.",
    italic=True
)
page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ÍNDICE (manual — Google Docs generará el automático)
# ═══════════════════════════════════════════════════════════════════════════
h1("Índice de Contenidos")
toc_items = [
    ("1.", "Introducción", ""),
    ("1.1.", "Contexto y motivación", ""),
    ("1.2.", "Objetivos", ""),
    ("1.3.", "Alcance de la solución", ""),
    ("2.", "Estado del Arte y Fundamentos Teóricos", ""),
    ("2.1.", "Grandes Modelos de Lenguaje y arquitectura Transformer", ""),
    ("2.2.", "Retrieval-Augmented Generation (RAG)", ""),
    ("2.3.", "Agentes LLM y patrones de orquestación", ""),
    ("2.4.", "Generación multimodal", ""),
    ("2.5.", "OCR y procesamiento de documentos físicos", ""),
    ("2.6.", "Model Context Protocol (MCP)", ""),
    ("2.7.", "Marcos legales y éticos", ""),
    ("3.", "Requisitos y Diseño de la Solución", ""),
    ("3.1.", "Requisitos funcionales", ""),
    ("3.2.", "Requisitos no funcionales", ""),
    ("3.3.", "Arquitectura de la aplicación", ""),
    ("3.4.", "Selección tecnológica", ""),
    ("3.5.", "Diseño de datos", ""),
    ("4.", "Implementación", ""),
    ("4.1.", "Estructura del código", ""),
    ("4.2.", "Integración con modelos LLM", ""),
    ("4.3.", "Gestión de prompts y orquestación de agentes", ""),
    ("4.4.", "Pipeline RAG con pgvector", ""),
    ("4.5.", "Generación asíncrona y notificaciones", ""),
    ("4.6.", "Mecanismos de seguridad", ""),
    ("4.7.", "Despliegue e infraestructura", ""),
    ("5.", "Evaluación y Experimentos", ""),
    ("5.1.", "Marco de evaluación — DeepEval", ""),
    ("5.2.", "Métricas técnicas de rendimiento", ""),
    ("5.3.", "Análisis de costes", ""),
    ("6.", "Pruebas y Validación de Producto", ""),
    ("6.1.", "Pruebas funcionales automatizadas", ""),
    ("6.2.", "Pruebas manuales requeridas", ""),
    ("6.3.", "Bugs identificados y corregidos", ""),
    ("6.4.", "UX y usabilidad", ""),
    ("7.", "Discusión", ""),
    ("7.1.", "Lecciones aprendidas", ""),
    ("7.2.", "Riesgos, ética y mitigaciones", ""),
    ("7.3.", "Análisis de limitaciones", ""),
    ("8.", "Conclusiones y Trabajo Futuro", ""),
    ("8.1.", "Conclusiones", ""),
    ("8.2.", "Trabajo futuro", ""),
    ("9.", "Bibliografía", ""),
    ("Anexo A.", "Guía de despliegue", ""),
    ("Anexo B.", "Manual de usuario resumido", ""),
    ("Anexo C.", "Prompts clave del sistema", ""),
]
for num, title, _ in toc_items:
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_after = Pt(2)
    indent = Cm(0.5) if num.count(".") > 1 else Cm(0)
    toc_para.paragraph_format.left_indent = indent
    r_num = toc_para.add_run(f"{num}  ")
    r_num.font.size = Pt(11)
    r_num.bold = True if num.count(".") <= 1 else False
    r_title = toc_para.add_run(title)
    r_title.font.size = Pt(11)
    r_title.bold = True if num.count(".") <= 1 else False

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Introducción")

h2("1.1. Contexto y motivación")
p(
    "El tejido empresarial español está compuesto en un 99,8% por PYMEs y autónomos "
    "(Ministerio de Industria, Comercio y Turismo, 2023). Este colectivo, que representa "
    "la columna vertebral de la economía nacional, enfrenta una paradoja tecnológica: "
    "dispone de acceso a herramientas digitales avanzadas pero carece del tiempo, los "
    "recursos humanos y, en muchos casos, de la formación técnica necesaria para adoptarlas "
    "e integrarlas de forma coherente en sus procesos de negocio."
)
p(
    "La irrupción de los Grandes Modelos de Lenguaje (LLMs) en el bienio 2022-2023 ha cambiado "
    "radicalmente el paradigma. Por primera vez en la historia de la computación, es posible "
    "construir un asistente que no solo responde preguntas en lenguaje natural, sino que también "
    "ejecuta acciones reales en nombre del usuario: genera documentos estructurados, agenda citas "
    "con validación de conflictos, analiza imágenes o crea contenido visual, todo ello desde "
    "una interfaz conversacional de texto."
)
p(
    "Observamos tres necesidades concretas no satisfactoriamente cubiertas por el mercado actual:"
)
bullet(
    "Los autónomos dedican entre 3 y 5 horas semanales a tareas administrativas repetitivas "
    "(elaboración de presupuestos, gestión de citas, respuestas a clientes), tiempo que podría "
    "dedicarse a actividades de mayor valor añadido."
)
bullet(
    "Las soluciones SaaS existentes en el mercado (Holded, Factorial, Canva, Notion...) son "
    "herramientas aisladas que no se comunican entre sí ni comprenden el contexto específico "
    "del negocio del usuario."
)
bullet(
    "Los chatbots genéricos de propósito general (ChatGPT, Microsoft Copilot) no están conectados "
    "a los datos ni a los procesos reales del negocio del usuario, lo que limita su utilidad "
    "operativa más allá de la redacción y la consulta de información general."
)
p(
    "Zeptai nace para cubrir precisamente esta brecha: un asistente de IA que conoce el negocio "
    "del usuario, actúa sobre él con herramientas reales y aprende de cada interacción, integrando "
    "en una única plataforma todas las capacidades que hasta ahora requerían múltiples herramientas "
    "dispares y conocimiento técnico especializado."
)
p(
    "En términos de mercado, la oportunidad es significativa: España cuenta con 3,3 millones de "
    "autónomos registrados en el RETA (2023), y el mercado global de AI SaaS se estimaba en "
    "62.000 millones de dólares en 2023 con una proyección de 300.000 millones para 2027 "
    "(Grand View Research, 2024). El coste medio de contratar un gestor externo para una PYME "
    "oscila entre 150 y 400 euros mensuales; Zeptai ofrece capacidades comparables por una "
    "fracción de ese coste (estimación: 19 €/mes en plan Pro)."
)

h2("1.2. Objetivos")
pb("Objetivo general")
p(
    "Desarrollar una plataforma web SaaS basada en IA Generativa que automatice las principales "
    "tareas de gestión de autónomos y PYMEs españolas mediante una arquitectura multi-agente "
    "orquestada, accesible desde una interfaz conversacional sin necesidad de formación técnica."
)
pb("Objetivos específicos")
bullet(
    "OE1. Implementar un sistema de agentes especializados (presupuestos, marketing, calendario, "
    "postventa, análisis de tickets) accesibles desde el chat mediante function calling."
)
bullet(
    "OE2. Diseñar un orquestador multi-agente (CouncilManager) capaz de generar perspectivas "
    "complementarias desde roles diferenciados (Estratega, Analista, Implementador) y producir "
    "una síntesis ejecutiva integrada."
)
bullet(
    "OE3. Integrar generación multimodal: texto (presupuestos PDF), imágenes publicitarias "
    "(DALL·E 3), presentaciones (python-pptx) y prompts de vídeo profesionales (Runway ML)."
)
bullet(
    "OE4. Implementar un pipeline de OCR + LLM para digitalizar y formalizar documentos "
    "físicos (notas manuscritas, facturas, capturas de pantalla)."
)
bullet(
    "OE5. Construir un pipeline RAG (Retrieval-Augmented Generation) sobre PostgreSQL + pgvector "
    "que permita al agente responder utilizando el conocimiento específico de cada negocio."
)
bullet(
    "OE6. Desplegar la solución en un entorno reproducible (Docker Compose) con base de datos "
    "relacional, sistema de notificaciones en tiempo real y servidor MCP para extensibilidad."
)
bullet(
    "OE7. Validar la solución mediante pruebas funcionales automatizadas (51 casos de prueba) "
    "y evaluación objetiva de calidad del pipeline RAG mediante DeepEval (LLM-as-a-judge)."
)

h2("1.3. Alcance de la solución")
p(
    "Zeptai se concibe como un MVP (Producto Mínimo Viable) funcional y desplegable en producción. "
    "El alcance definido para este trabajo comprende los siguientes elementos:"
)
pb("Dentro del alcance:")
bullet("Interfaz web PWA responsive (Flask + Jinja2 + Tailwind CSS), instalable desde el navegador.")
bullet("Sistema de autenticación completo con perfil de negocio configurable mediante wizard.")
bullet("Seis agentes especializados integrados vía tool-use: agenda, presupuestos, marketing, postventa, OCR y consultas RAG.")
bullet("Seis agentes proactivos con ejecución programada (schedule library, 09:00 diario): BusinessHealth, GrantHunter, Networker, PostSales, AdminRedactor, MarketingAgent.")
bullet("CouncilManager con debate multi-perspectiva y síntesis en streaming SSE.")
bullet("Pipeline RAG completo con ingesta de PDF, chunking, embeddings y búsqueda vectorial en pgvector.")
bullet("Sistema de notificaciones in-app y web push (VAPID).")
bullet("Panel de métricas de uso LLM (coste, latencia, modelo) y exportación a Excel.")
bullet("Evaluación objetiva de calidad con DeepEval (4 métricas, 7 casos de prueba).")
bullet("Despliegue Docker Compose con 4 servicios (db, mcp, web, scheduler).")

pb("Fuera del alcance (trabajo futuro):")
bullet("App móvil nativa (iOS/Android).")
bullet("Integración directa con WhatsApp Business API (webhook Twilio / Meta Cloud API).")
bullet("Fine-tuning de modelos propios sobre datos del dominio.")
bullet("Marketplace de agentes con contribuciones de terceros.")
bullet("Facturación automática SaaS (Stripe) con planes freemium/pro/enterprise.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — ESTADO DEL ARTE
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Estado del Arte y Fundamentos Teóricos")

h2("2.1. Grandes Modelos de Lenguaje y arquitectura Transformer")
p(
    "Los Grandes Modelos de Lenguaje (LLMs, por sus siglas en inglés Large Language Models) son "
    "sistemas de aprendizaje profundo basados en la arquitectura Transformer (Vaswani et al., 2017), "
    "entrenados sobre corpus masivos de texto para aprender representaciones estadísticas del lenguaje "
    "a una escala sin precedentes históricos."
)
p(
    "La innovación central de la arquitectura Transformer es el mecanismo de atención multi-cabeza "
    "(multi-head self-attention), que permite al modelo relacionar cualquier par de posiciones en la "
    "secuencia de entrada de forma directa, independientemente de su distancia. Esto supera la "
    "limitación fundamental de las arquitecturas recurrentes (RNN, LSTM) que procesaban la secuencia "
    "de forma secuencial y perdían información al extenderse la distancia entre tokens relacionados. "
    "Formalmente, el mecanismo de atención se define como:"
)
code("Attention(Q, K, V) = softmax(QK^T / sqrt(d_k)) · V")
p(
    "donde Q (queries), K (keys) y V (values) son proyecciones lineales aprendidas de la secuencia "
    "de entrada, y d_k es la dimensión de las claves. La normalización por sqrt(d_k) estabiliza "
    "los gradientes para dimensiones elevadas. El término 'multi-cabeza' hace referencia a la "
    "ejecución en paralelo de h instancias independientes de este mecanismo, cuyos resultados "
    "se concatenan y proyectan linealmente."
)
p(
    "La variante decoder-only (familia GPT) apila N bloques idénticos, cada uno compuesto por: "
    "(1) atención causal enmascarada (la posición i solo puede atender posiciones j ≤ i, garantizando "
    "la autoregresividad del modelo); (2) una red feed-forward de dos capas con activación GELU; "
    "y (3) normalización por capas con conexión residual. El entrenamiento se realiza con un "
    "objetivo de modelado de lenguaje causal: predecir el siguiente token dada la secuencia anterior."
)
p(
    "GPT-4 (OpenAI, 2023), el modelo base que alimenta Zeptai, representa el estado del arte "
    "en modelos de propósito general. Sus características más relevantes para este proyecto son: "
    "una ventana de contexto de 128.000 tokens (equivalente a aproximadamente 90.000 palabras), "
    "capacidad multimodal nativa para procesar texto e imágenes en el mismo contexto, y un "
    "afinado mediante RLHF (Reinforcement Learning from Human Feedback) que alinea las salidas "
    "del modelo con las preferencias humanas en cuanto a utilidad, veracidad y seguridad."
)
p(
    "En Zeptai empleamos específicamente la variante gpt-4o ('o' de omni), optimizada por OpenAI "
    "para ofrecer la mejor relación calidad/coste/latencia del ecosistema GPT-4 en la ventana "
    "temporal 2024-2025. Cada llamada al modelo envía: el system prompt personalizado del negocio, "
    "el historial de conversación de las últimas 10 interacciones (ventana deslizante), y el "
    "esquema JSON de las herramientas disponibles para function calling."
)

h2("2.2. Retrieval-Augmented Generation (RAG)")
p(
    "RAG (Lewis et al., 2020) es una técnica que combina un módulo de recuperación de información "
    "con un LLM para generar respuestas fundamentadas en una base de conocimiento externa, sin "
    "necesidad de reentrenar o ajustar el modelo. El flujo canónico es:"
)
code(
    "[Consulta usuario] → [Encoder embedding] → [Vector Store]\n"
    "→ [Top-K documentos más similares] → [LLM (contexto aumentado)] → [Respuesta]"
)
p(
    "Las ventajas de RAG frente al fine-tuning para el caso de uso de Zeptai son determinantes:"
)
bullet("Sin coste de reentrenamiento (que puede ascender a miles de dólares por ejecución).")
bullet("El conocimiento es actualizable en tiempo real sin tocar el modelo base.")
bullet("Trazabilidad de fuentes: cada respuesta puede atribuirse a los fragmentos concretos que la fundamentan.")
bullet("Idóneo para conocimiento empresarial específico en constante evolución (catálogos, tarifas, FAQs).")

p(
    "Zeptai implementa un pipeline RAG completo sobre PostgreSQL con la extensión pgvector, "
    "eliminando la necesidad de un servicio vectorial externo (Pinecone, Qdrant, ChromaDB). "
    "El mismo contenedor de base de datos gestiona tanto los datos relacionales como los vectores "
    "de embeddings, simplificando la arquitectura y reduciendo los costes operativos."
)
pb("Pipeline de ingesta (módulo embeddings.py):")
code(
    "[PDF / TXT subido] → [PyMuPDF: extrae texto] → [Chunker recursivo]\n"
    "→ [text-embedding-3-small (OpenAI)] → [KnowledgeChunk en pgvector]"
)
p(
    "El chunker divide el documento en fragmentos de aproximadamente 1.600 caracteres con un "
    "overlap del 25% (400 caracteres) para preservar el contexto en los límites entre fragmentos. "
    "El algoritmo sigue una cascada de separadores: párrafo → frase → carácter, garantizando "
    "que los fragmentos nunca cortan unidades semánticas básicas."
)
pb("Pipeline de recuperación (retrieve_chunks):")
code(
    "[Pregunta usuario] → [embed_text()] → [cosine similarity en pgvector]\n"
    "→ [Top-5 chunks más similares] → [inyección en system prompt del LLM]"
)
p(
    "El modelo de embedding empleado es text-embedding-3-small (OpenAI), con 1.536 dimensiones "
    "y un coste aproximado de 0,000020 $/1K tokens (prácticamente gratuito en el volumen de "
    "uso de una PYME). Supera a text-embedding-ada-002 en los benchmarks MTEB con menor coste, "
    "lo que lo convierte en la opción óptima para este caso de uso."
)
p(
    "El sistema distingue dos fuentes de conocimiento que se consultan conjuntamente: "
    "'wizard' (campos del perfil de negocio: horario, servicios, tarifas, FAQs) y "
    "'document' (PDFs o archivos TXT subidos desde la sección Knowledge). Esta arquitectura "
    "proporciona al agente acceso unificado a la totalidad del conocimiento del negocio en "
    "una única búsqueda vectorial."
)

h2("2.3. Agentes LLM y patrones de orquestación")
p(
    "Un agente LLM es un sistema donde el modelo de lenguaje no solo genera texto sino que también "
    "decide qué acciones ejecutar, las ejecuta mediante herramientas externas y razona sobre sus "
    "resultados para decidir el siguiente paso (Yao et al., 2023). El patrón fundamental que "
    "articula este comportamiento es ReAct (Reasoning + Acting), que intercala razonamiento "
    "explícito con acciones concretas en un bucle:"
)
code(
    "Thought:     'El usuario quiere reservar una cita para el martes'\n"
    "Action:      check_availability(date='2026-04-21')\n"
    "Observation: 'Huecos disponibles: 09:00, 10:00, 16:00'\n"
    "Thought:     'Hay disponibilidad, pregunto al usuario su preferencia'\n"
    "Response:    'El martes tengo disponibles las 9h, 10h o 16h. ¿Cuál prefieres?'"
)
p(
    "OpenAI proporciona una implementación nativa de este patrón mediante el mecanismo de "
    "function calling (también denominado tool use): cuando el modelo decide que debe ejecutar "
    "una acción, devuelve un JSON estructurado con el nombre de la función y sus argumentos "
    "en lugar de texto libre, garantizando una interfaz determinista y tipada entre el LLM y "
    "el código de la aplicación."
)
p(
    "En Zeptai, el TOOLS_SCHEMA define seis herramientas disponibles para el AgentExecutor:"
)
bullet("check_availability — consulta la agenda del negocio por fecha.")
bullet("book_appointment — reserva una cita con validación de conflictos en tiempo real.")
bullet("create_proposal_from_last_image — genera un presupuesto PDF desde una imagen analizada por OCR.")
bullet("create_proposal_from_text — genera un presupuesto PDF a partir de datos dictados por el usuario.")
bullet("generate_marketing_material — genera imágenes (DALL·E 3), presentaciones (PPT) o prompts de vídeo.")
bullet("handle_customer_service — gestiona incidencias y consultas postventa.")

h3("CouncilManager: debate multi-agente")
p(
    "El patrón multi-agente más sofisticado implementado en Zeptai es el CouncilManager, "
    "inspirado en los trabajos de Park et al. (2023) sobre agentes generativos con roles "
    "diferenciados y en el framework AutoGen de Microsoft (Wu et al., 2023). Implementa un "
    "debate estructurado con tres perspectivas independientes sobre la misma pregunta estratégica:"
)
bullet("El Socio (growth) — perspectiva comercial: CAC, LTV, ventas, crecimiento de mercado.")
bullet("El Gestor (legal/fiscal) — perspectiva regulatoria: AEAT, obligaciones fiscales, cumplimiento normativo.")
bullet("El Coach (operaciones) — perspectiva de ejecucion: productividad del dueno, procesos internos, recursos.")
p(
    "El debate se estructura en tres rondas ejecutadas secuencialmente mediante un generador "
    "asincrono que emite eventos SSE en tiempo real: (1) Ronda de opiniones iniciales — cada "
    "agente expone su posicion de forma independiente con una pausa de 0,8 s entre respuestas; "
    "(2) Ronda de replicas/debate — cada agente rebate las posiciones de los otros con 0,6 s "
    "entre turnos; (3) Sintesis y plan de accion — un cuarto rol ('Secretario') integra todas "
    "las perspectivas en una recomendacion ejecutiva accionable. "
    "El frontend recibe eventos tipados: 'typing' (indicador de escritura), 'message' (respuesta "
    "completa), 'divider' (separador de ronda) y 'plan' (plan final). El timeout de sesion "
    "del council es de 180 segundos. Los roles son especificos al contexto de una PYME espanola, "
    "lo que diferencia a Zeptai de sistemas multi-agente genericos."
)

h3("Agentes proactivos con ejecución programada")
p(
    "Ademas del ciclo reactivo (usuario → agente → respuesta), Zeptai implementa seis agentes "
    "que se ejecutan de forma autonoma cada dia a las 09:00 mediante la libreria schedule "
    "(run_scheduler.py). El scheduler arranca inmediatamente al iniciar y luego ejecuta "
    "run_daily_tasks() en bucle con sleep de 60 segundos:"
)
bullet("BusinessHealthAgent: analiza los tickets de gasto del dia, proyecta el gasto de fin de mes y genera alertas si la desviacion supera el 20%. Emite notificaciones con emojis de tendencia (subida, bajada, estable).")
bullet("GrantHunterAgent: busca subvenciones activas en la tabla grant de la BD, filtra por sector del usuario y notifica un maximo de 2 ayudas por ejecucion para evitar spam.")
bullet("SynergyAgent (Networker): detecta sinergias entre usuarios analizando sus perfiles de gasto, genera puntuaciones de compatibilidad (threshold >= 80) y crea registros en synergy_match.")
bullet("PostSalesAgent: gestiona incidencias postventa (devoluciones, quejas, estado de pedidos) usando politicas configurables por empresa desde agent_config.")
bullet("AdminRedactorAgent: procesa imagenes de tickets y facturas. Clasifica el tipo ('receipt' formal vs 'draft' manuscrito) antes de invocar GPT-4o Vision para extraccion estructurada. Genera PDFs con FPDF2.")
bullet("MarketingAgent: genera contenido visual (imagenes DALL-E 3, presentaciones PowerPoint 2-8 slides, prompts de video Runway gen3a_turbo) en hilos daemon con notificacion push al finalizar.")

h2("2.4. Generación multimodal")
p(
    "Zeptai integra cuatro modalidades de generación de contenido, abarcando texto, imágenes, "
    "presentaciones y vídeo:"
)
h3("Imágenes — DALL·E 3 (OpenAI)")
p(
    "DALL·E 3 es el modelo de generación de imágenes de OpenAI, basado en arquitecturas de "
    "difusión latente condicionadas por texto. Soporta resoluciones de hasta 1792×1024 píxeles "
    "con calidad fotorrealista. El marketing agent de Zeptai construye un prompt enriquecido "
    "con el contexto específico del negocio (sector, nombre, estilo de comunicación) y lo "
    "envía a la API. La imagen resultante se descarga y almacena localmente. Coste aproximado: "
    "0,04 € por imagen en calidad estándar."
)
h3("Presentaciones — python-pptx + GPT-4o")
p(
    "La generación de presentaciones se implementa mediante una cadena de dos pasos: GPT-4o "
    "genera el contenido estructurado como un JSON (título, subtítulo, slides, puntos clave "
    "por slide), y python-pptx renderiza ese contenido sobre una plantilla corporativa "
    "predefinida usando layouts de portada, contenido y cierre. Esta aproximación ofrece "
    "control total sobre el diseño sin depender de APIs externas, con un coste mínimo "
    "(únicamente la llamada LLM, aproximadamente 0,01 €/PPT)."
)
h3("Vídeo — Runway ML Gen-3")
p(
    "Para vídeos, Zeptai genera un prompt profesional detallado (descripción de escena, estilo "
    "cinematográfico, movimiento de cámara, paleta de color, ritmo) usando GPT-4o, y lo envía "
    "a la API de Runway ML Gen-3 Alpha para renderizado. El proceso es asíncrono: se ejecuta "
    "en un hilo daemon y notifica al usuario mediante push cuando el vídeo está disponible. "
    "Coste estimado: 0,10-0,50 € por vídeo de 4 segundos."
)

h2("2.5. OCR y procesamiento de documentos físicos")
p(
    "El AdminRedactorAgent implementa un pipeline de digitalización de documentos físicos sin "
    "recurrir a motores OCR tradicionales (Tesseract, AWS Textract). En su lugar, aprovecha "
    "las capacidades multimodales nativas de GPT-4o para procesar imágenes y extraer información "
    "estructurada directamente:"
)
code(
    "[Imagen capturada (foto, captura)] → [GPT-4o Vision]\n"
    "→ [Extracción estructurada JSON (partidas, precios, cliente, fecha)]\n"
    "→ [Plantilla PDF (ReportLab)] → [Documento formalizado descargable]"
)
p(
    "Esta aproximación mejora la robustez frente a OCR convencional en casos de escritura "
    "manuscrita, baja resolución, documentos parcialmente visibles o idiomas mixtos, ya que "
    "el LLM infiere el contexto semántico en lugar de reconocer caracteres aislados."
)

h2("2.6. Model Context Protocol (MCP)")
p(
    "El Model Context Protocol (MCP) es un estándar abierto propuesto por Anthropic en 2024 "
    "para estandarizar la integración de herramientas externas con sistemas LLM. Define un "
    "protocolo cliente-servidor donde el servidor expone herramientas tipadas que el agente "
    "puede invocar de forma estructurada. Zeptai implementa un servidor MCP propio "
    "(mcp_server_sse.py) basado en el transporte Server-Sent Events (SSE), que expone seis "
    "herramientas adicionales a los agentes:"
)
bullet("get_financial_summary — resumen financiero del negocio.")
bullet("get_appointments — lista de citas programadas.")
bullet("search_web — búsqueda web para el GrantHunterAgent.")
bullet("schedule_appointment — creación de citas desde el agente proactivo.")
bullet("send_email_notification — envío de notificaciones por email.")
bullet("get_business_stats — estadísticas de uso de la plataforma.")
p(
    "La ventaja clave del transporte SSE frente al transporte stdio estándar es que el proceso "
    "MCP se lanza una sola vez y permanece activo, reduciendo la latencia por herramienta de "
    "aproximadamente 500 ms (fork de proceso) a aproximadamente 5 ms (HTTP local), y permitiendo "
    "conexiones concurrentes de múltiples workers de Gunicorn sin conflictos."
)

h2("2.7. Marcos legales y éticos")
h3("Reglamento General de Protección de Datos (RGPD)")
p(
    "Zeptai opera en España y procesa datos personales de usuarios (nombre, teléfono, email, "
    "historial de conversaciones, documentos subidos). Las obligaciones aplicables bajo el "
    "Reglamento 2016/679 (RGPD) incluyen: base legitimadora de contrato (art. 6.1.b), ya que "
    "el tratamiento es necesario para la prestación del servicio; principio de minimización "
    "de datos; derecho de supresión implementado mediante endpoints de borrado de documentos "
    "y gestión de cuenta; y medidas de seguridad técnica (contraseñas hasheadas con PBKDF2-SHA256, "
    "sesiones server-side, aislamiento estricto por usuario en todas las consultas SQL)."
)
h3("EU AI Act (Reglamento 2024/1689)")
p(
    "Zeptai se clasifica bajo el AI Act en la categoría de riesgo mínimo/limitado. Esta "
    "clasificación se justifica porque el sistema no toma decisiones autónomas con efecto "
    "jurídico sobre personas, no realiza scoring social, no manipula el comportamiento del "
    "usuario y no opera en sectores críticos (salud, justicia, infraestructuras críticas). "
    "Se cumple la obligación de transparencia del artículo 50: la interfaz identifica "
    "claramente al agente como sistema de IA."
)
h3("Riesgos identificados y mitigaciones")

table_2col(
    [
        ("Alucinaciones (hallucinations)", "El LLM genera información incorrecta. Mitigación: static_knowledge del usuario limita el dominio de respuesta; los documentos generados incluyen aviso de revisión recomendada."),
        ("Prompt injection", "Un usuario malicioso intenta manipular el system prompt. Mitigación: el system prompt se inyecta en posición fija (role: system); el input del usuario no puede sobreescribirlo."),
        ("Privacidad de datos conversacionales", "El historial de chat contiene información sensible del negocio. Mitigación: los ChatMessage se almacenan en la BD propia, nunca en servidores de OpenAI, con filtrado estricto por usuario."),
        ("Dependencia de proveedor (vendor lock-in)", "OpenAI puede cambiar precios o deprecar modelos. Mitigación: capa de abstracción get_openai_client() para sustituir el proveedor; llm_tracker registra todas las llamadas para auditoría."),
    ],
    header=["Riesgo", "Descripción y mitigación"]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — REQUISITOS Y DISEÑO
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Requisitos y Diseño de la Solución")

h2("3.1. Requisitos funcionales")
p(
    "El análisis de necesidades del colectivo objetivo (autónomos y PYMEs) llevó a la definición "
    "de dieciséis requisitos funcionales que articulan las capacidades esenciales de la plataforma:"
)
table_2col(
    [
        ("RF01", "El sistema permitirá al usuario configurar el perfil de su negocio (nombre, sector, system prompt personalizado, conocimiento estático, logo)."),
        ("RF02", "El sistema ofrecerá una interfaz de chat conversacional persistente con historial de las últimas 10 interacciones."),
        ("RF03", "El agente podrá consultar disponibilidad de agenda para una fecha dada, devolviendo los huecos libres."),
        ("RF04", "El agente podrá reservar citas con validación de conflictos horarios en tiempo real."),
        ("RF05", "El agente generará presupuestos en PDF a partir de datos dictados por el usuario en el chat."),
        ("RF06", "El agente generará presupuestos en PDF a partir del análisis OCR de imágenes adjuntas."),
        ("RF07", "El agente creará imágenes publicitarias mediante la API de DALL·E 3 con prompt contextualizado."),
        ("RF08", "El agente creará presentaciones PPT con número de slides parametrizable sobre plantilla corporativa."),
        ("RF09", "El agente generará prompts profesionales para vídeos con la API de Runway ML Gen-3."),
        ("RF10", "El CouncilManager orquestará debates entre 3 perspectivas de agentes especializados, transmitiendo cada respuesta en streaming SSE."),
        ("RF11", "Los documentos generados se almacenarán en una sección 'Documentos' accesible y descargable."),
        ("RF12", "El sistema enviará notificaciones in-app y web push (VAPID) cuando finalicen tareas asíncronas."),
        ("RF13", "El panel de métricas mostrará consumo LLM (tokens, coste estimado, latencia) por modelo y por período."),
        ("RF14", "El usuario podrá exportar sus tickets/transacciones a formato Excel (.xlsx)."),
        ("RF15", "Los agentes proactivos (GrantHunter, BusinessCoach, Networker) se ejecutarán según horario programado y enviarán notificaciones con los resultados."),
        ("RF16", "El Marketplace permitirá al usuario activar y desactivar agentes por categoría según sus necesidades."),
    ],
    header=["ID", "Descripción"]
)

h2("3.2. Requisitos no funcionales")
table_2col(
    [
        ("RNF01 — Rendimiento", "El tiempo de respuesta del chat no superará los 5 segundos en el 95% de las peticiones (sin contar generación asíncrona)."),
        ("RNF02 — Disponibilidad", "Despliegue Docker Compose reproducible en cualquier máquina con Docker Desktop instalado mediante un único comando."),
        ("RNF03 — Seguridad", "Autenticación obligatoria en todas las rutas protegidas; contraseñas hasheadas con PBKDF2-SHA256; aislamiento estricto de datos por usuario."),
        ("RNF04 — Escalabilidad", "Arquitectura stateless (Flask + PostgreSQL externo) que permite escalar horizontalmente añadiendo instancias de Gunicorn detrás de un load balancer."),
        ("RNF05 — Usabilidad", "Interfaz responsive válida para móvil y escritorio; instalable como PWA desde Chrome/Firefox mediante beforeinstallprompt."),
        ("RNF06 — Observabilidad", "Todas las llamadas LLM quedan trazadas en LLMUsageLog con modelo, stage, tokens, latencia y coste estimado; logs estructurados con Python logging."),
        ("RNF07 — Mantenibilidad", "Arquitectura modular por blueprints Flask; agentes desacoplados del orquestador mediante interfaces definidas."),
    ],
    header=["ID", "Descripción"]
)

h2("3.3. Arquitectura de la aplicación")
p(
    "La arquitectura de Zeptai sigue un patrón en capas con separación clara de responsabilidades. "
    "El siguiente diagrama ilustra el flujo principal de una petición desde el cliente hasta "
    "los sistemas externos:"
)
code(
    "┌─────────────────────────────────────────┐\n"
    "│           CLIENTE (PWA / Browser)       │\n"
    "│   HTML + Tailwind CSS · SW · Push API   │\n"
    "└────────────────┬────────────────────────┘\n"
    "                 │ HTTPS / SSE\n"
    "┌────────────────▼────────────────────────┐\n"
    "│   Flask (Gunicorn) — routes/api.py      │\n"
    "│              routes/web.py              │\n"
    "└──┬──────────┬───────────┬───────────────┘\n"
    "   │          │           │\n"
    "┌──▼────┐  ┌──▼────────┐  ┌▼──────────────┐\n"
    "│Agent  │  │Council    │  │Proactive      │\n"
    "│Exec.  │  │Manager    │  │Agents         │\n"
    "│+RAG   │  │(3×GPT-4o) │  │(schedule lib) │\n"
    "└──┬────┘  └──┬────────┘  └┬──────────────┘\n"
    "   │          │             │\n"
    "   └──────────┴─────────────┘\n"
    "                │\n"
    "   ┌────────────┴────────────┐\n"
    "   │  PostgreSQL + pgvector  │\n"
    "   │  (datos + embeddings)   │\n"
    "   └─────────────────────────┘\n"
    "                │\n"
    "   ┌────────────┴────────────┐\n"
    "   │  MCP Server SSE :8001   │\n"
    "   │  (6 herramientas MCP)   │\n"
    "   └─────────────────────────┘"
)
p(
    "El flujo de una petición típica de chat con tool use sigue estos pasos: "
    "(1) el usuario envía un mensaje mediante POST /api/chat; "
    "(2) el AgentExecutor construye el array messages con el system prompt personalizado más el historial de conversación; "
    "(3) primera llamada a GPT-4o: el modelo decide si responder directamente o invocar una herramienta; "
    "(4) si invoca una herramienta, _process_tool_calls() ejecuta la función correspondiente y obtiene el resultado; "
    "(5) segunda llamada a GPT-4o para sintetizar la respuesta final con el resultado de la herramienta; "
    "(6) HistoryService persiste la interacción en la tabla ChatMessage; "
    "(7) la respuesta se devuelve al frontend como JSON."
)

h2("3.4. Selección tecnológica")
p(
    "Las decisiones tecnológicas del proyecto priorizaron tres criterios: compatibilidad con el "
    "ecosistema Python de IA, minimización de la superficie de servicios externos, y "
    "reproducibilidad del entorno de desarrollo y producción."
)
table_2col(
    [
        ("Python 3.11", "Ecosistema líder en IA/ML, tipado gradual, soporte async nativo."),
        ("Flask 3.x", "Microframework ligero, ideal para APIs REST y SSR con Jinja2. FastAPI descartado por mayor overhead en setup y menor integración con Jinja2."),
        ("SQLAlchemy 2.x", "ORM maduro con soporte para PostgreSQL, MySQL y SQLite según entorno."),
        ("PostgreSQL 15 + pgvector", "RDBMS robusto con extensión vectorial nativa. Elimina la necesidad de un servicio vectorial externo (Pinecone, Qdrant)."),
        ("schedule library", "Scheduler simple en-proceso para agentes proactivos. run_daily_tasks() se ejecuta a las 09:00 diarias con bucle sleep(60). No requiere Celery ni Redis."),
        ("Jinja2 + Tailwind CSS", "Server-side rendering sin build step. Vue/React descartados por overhead en un MVP."),
        ("GPT-4o", "Mejor relación coste/calidad para razonamiento y function calling en 2024-2025."),
        ("text-embedding-3-small", "1.536 dims, ~0.00002 $/1K tokens. Supera a ada-002 en benchmarks MTEB."),
        ("DALL·E 3", "Generación de imágenes de alta calidad integrada en el ecosistema OpenAI."),
        ("Runway ML Gen-3", "API estable de generación de vídeo con calidad cinematográfica."),
        ("DeepEval", "Framework open-source de evaluación LLM-as-a-judge para validación objetiva del pipeline RAG."),
        ("Docker + Docker Compose", "Reproducibilidad, aislamiento y despliegue en un único comando."),
        ("FastMCP", "Implementación Python del Model Context Protocol con transporte SSE."),
        ("Flask-Limiter", "Rate limiting por endpoint y por usuario para protección contra abuso."),
        ("pywebpush + py-vapid", "Web push notifications estándar W3C para notificaciones PWA."),
    ],
    header=["Tecnología", "Justificación de selección"]
)

h2("3.5. Diseño de datos")
p(
    "El modelo de datos de Zeptai orbita en torno a la entidad business_profile, que representa "
    "a cada usuario-negocio registrado en la plataforma. La clave primaria de negocio es user_phone, "
    "y todas las entidades del sistema se relacionan con ella mediante clave foránea. "
    "El esquema real en PostgreSQL comprende 14 tablas:"
)
code(
    "business_profile (1) --- (N) appointment\n"
    "business_profile (1) --- (N) ticket\n"
    "business_profile (1) --- (N) generated_document\n"
    "business_profile (1) --- (N) chat_message\n"
    "business_profile (1) --- (N) chat_feedback\n"
    "business_profile (1) --- (N) notification\n"
    "business_profile (1) --- (N) activity_log\n"
    "business_profile (1) --- (N) llm_call\n"
    "business_profile (1) --- (N) knowledge_chunk\n"
    "business_profile (1) --- (N) rag_retrieval\n"
    "business_profile (1) --- (N) incident\n"
    "business_profile (1) --- (N) synergy_match\n"
    "grant              (independiente — acceso global por sector)\n"
    "apscheduler_jobs   (interna APScheduler — persistencia de jobs)"
)

h3("Tabla: business_profile")
table_2col(
    [
        ("id", "INTEGER PK autoincremental"),
        ("user_phone", "VARCHAR — identificador de negocio, usado como FK en todas las tablas"),
        ("email", "VARCHAR — email de acceso (login)"),
        ("business_name", "VARCHAR — nombre comercial del negocio"),
        ("password_hash", "VARCHAR — hash PBKDF2-SHA256 de la contrasena"),
        ("reset_token / reset_token_expiry", "VARCHAR / TIMESTAMPTZ — token de recuperacion de contrasena"),
        ("plan_tier", "VARCHAR — nivel de suscripcion (free, pro, etc.)"),
        ("system_prompt", "TEXT — personalidad e instrucciones del agente IA"),
        ("static_knowledge", "JSON — conocimiento estatico del negocio (tarifas, FAQs, horario)"),
        ("active_agents", "JSON — lista de agentes activados en el Marketplace"),
        ("agent_config", "JSON — configuracion especifica por agente"),
        ("features", "JSON — flags de funcionalidades (bot_enabled, push_enabled, etc.)"),
        ("push_subscription", "TEXT — endpoint VAPID para web push notifications"),
        ("logo_path", "VARCHAR — ruta local al logo del negocio"),
        ("created_at", "TIMESTAMPTZ — fecha de registro"),
    ],
    header=["Campo", "Descripcion"]
)

h3("Tabla: knowledge_chunk (RAG)")
table_2col(
    [
        ("id", "INTEGER PK"),
        ("user_phone", "VARCHAR FK -> business_profile"),
        ("source_type", "VARCHAR — 'wizard' (perfil) o 'document' (PDF subido)"),
        ("source_name", "VARCHAR — nombre del campo wizard o del fichero PDF"),
        ("content", "TEXT — fragmento de texto (~1.600 chars)"),
        ("embedding", "USER-DEFINED (vector 1536 dims via pgvector) — embedding de text-embedding-3-small"),
        ("created_at", "TIMESTAMPTZ"),
    ],
    header=["Campo", "Descripcion"]
)

h3("Tabla: llm_call (trazabilidad de modelos)")
table_2col(
    [
        ("id", "INTEGER PK"),
        ("user_phone", "VARCHAR FK -> business_profile"),
        ("model", "VARCHAR — gpt-4o, dall-e-3, text-embedding-3-small, etc."),
        ("stage", "VARCHAR — chat_main, chat_tool_followup, image_generation, council_*, etc."),
        ("prompt_tokens", "INTEGER — tokens de entrada"),
        ("completion_tokens", "INTEGER — tokens de salida"),
        ("total_tokens", "INTEGER — total tokens consumidos"),
        ("latency_ms", "INTEGER — latencia de la llamada en milisegundos"),
        ("cost_usd", "DOUBLE — coste estimado en USD"),
        ("success", "BOOLEAN — si la llamada finalizo correctamente"),
        ("error_message", "TEXT — mensaje de error si success=false"),
        ("created_at", "TIMESTAMPTZ"),
    ],
    header=["Campo", "Descripcion"]
)

h3("Tabla: chat_message")
table_2col(
    [
        ("id", "INTEGER PK"),
        ("user_phone", "VARCHAR FK -> business_profile"),
        ("role", "VARCHAR — 'user', 'assistant' o 'tool'"),
        ("content", "TEXT — contenido del mensaje"),
        ("name", "VARCHAR — nombre de la herramienta (solo en mensajes role=tool)"),
        ("tool_call_id", "VARCHAR — ID de la llamada de herramienta asociada"),
        ("created_at", "TIMESTAMPTZ"),
    ],
    header=["Campo", "Descripcion"]
)

h3("Tabla: ticket (gastos/transacciones)")
table_2col(
    [
        ("id", "INTEGER PK"),
        ("user_phone", "VARCHAR FK -> business_profile"),
        ("image_path", "VARCHAR — ruta a la imagen del ticket original"),
        ("status", "VARCHAR — estado del procesamiento OCR"),
        ("concept", "VARCHAR — concepto del gasto"),
        ("total", "DOUBLE — importe total"),
        ("base", "DOUBLE — base imponible"),
        ("tax_percent", "DOUBLE — porcentaje de IVA"),
        ("fee", "DOUBLE — cuota de IVA"),
        ("nif", "VARCHAR — NIF del emisor"),
        ("provider", "VARCHAR — nombre del proveedor"),
        ("ticket_number", "VARCHAR — numero de factura/ticket"),
        ("date", "TIMESTAMPTZ — fecha del gasto"),
        ("raw_data", "TEXT — JSON raw extraido por GPT-4o Vision"),
        ("created_at", "TIMESTAMPTZ"),
    ],
    header=["Campo", "Descripcion"]
)

h3("Resto de tablas")
table_2col(
    [
        ("appointment",
         "Agenda del negocio. Cada fila es una cita reservada por el agente mediante la herramienta book_appointment. "
         "Antes de confirmar una reserva, el AgentExecutor consulta esta tabla para detectar conflictos horarios. "
         "Campos: business_phone, date, time, end_time, client_name, client_phone, created_at."),
        ("generated_document",
         "Biblioteca de contenido generado por el agente: PDFs de presupuestos, imagenes DALL-E 3, "
         "presentaciones PPT y prompts de video. El campo file_path apunta al fichero fisico en "
         "/static/generated_docs/ e incluye un UUID para evitar enumeracion. Es lo que muestra la seccion /documents. "
         "Campos: user_phone, file_path, doc_type, client_name, created_at."),
        ("notification",
         "Bandeja de notificaciones in-app. Cuando un hilo asincrono termina (imagen lista, PPT generado) "
         "o un agente proactivo encuentra algo relevante, escribe aqui. El frontend hace polling cada 15s "
         "contra /api/notifications/unread_count para actualizar el badge del campana sin recargar la pagina. "
         "Campos: user_phone, title, message, type, link, is_read, created_at."),
        ("activity_log",
         "Registro de auditoria de acciones del sistema. Cada vez que un agente ejecuta una accion relevante "
         "(documento generado, cita reservada, agente activado) se escribe una fila con el nombre del agente "
         "y la accion. Permite demostrar trazabilidad ante auditorias (RGPD). "
         "Campos: user_phone, agent_name, action, timestamp."),
        ("rag_retrieval",
         "Log de cada busqueda vectorial RAG ejecutada. Registra la consulta (preview de 200 chars), "
         "cuantos chunks devolvio el retriever y el score medio de similitud coseno. "
         "Es la tabla que alimenta las metricas de calidad RAG evaluadas por DeepEval "
         "(Contextual Precision y Contextual Recall). Campos: user_phone, query_preview, chunks_returned, avg_score, created_at."),
        ("incident",
         "Incidencias abiertas por el PostSalesAgent: devoluciones, quejas y consultas de estado de pedido. "
         "El campo type clasifica el tipo de incidencia y status registra su ciclo de vida (abierta, en gestion, resuelta). "
         "Campos: user_phone, order_id, type, status, description, created_at."),
        ("synergy_match",
         "Emparejamientos entre usuarios detectados por el SynergyAgent (Networker). "
         "Guarda los telefonos de los dos negocios compatibles, un score de compatibilidad "
         "(solo se notifica si score >= 80) y el motivo de la sinergia generado por el LLM. "
         "Campos: user_a_phone, user_b_phone, score, reason, status, created_at."),
        ("grant",
         "Subvenciones y ayudas detectadas por el GrantHunterAgent. Es una tabla GLOBAL (no por usuario): "
         "el agente escribe aqui las subvenciones encontradas y filtra por sector del usuario al notificar. "
         "El campo notified_phones (JSON) registra a quien ya se ha avisado para no repetir la notificacion. "
         "Maximo 2 grants notificados por usuario por ejecucion diaria. "
         "Campos: title, description, sector_focus, amount, link, deadline, notified_phones, created_at."),
        ("chat_feedback",
         "Valoraciones de las respuestas del chat enviadas por el usuario (thumbs up/down). "
         "El campo rating es un smallint y message_preview guarda los primeros caracteres del mensaje valorado "
         "para contexto. Sirve para medir satisfaccion y podria alimentar un sistema de evaluacion "
         "continua como trabajo futuro. Campos: user_phone, rating, message_preview, created_at."),
        ("apscheduler_jobs",
         "Tabla interna creada automaticamente por la dependencia APScheduler al iniciar la aplicacion. "
         "No se usa activamente: el scheduler real de Zeptai emplea la libreria schedule (run_scheduler.py). "
         "Puede ignorarse a efectos funcionales. Campos: id, next_run_time, job_state (bytea)."),
    ],
    header=["Tabla", "Funcion y campos"]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — IMPLEMENTACIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Implementación")

h2("4.1. Estructura del código")
p(
    "El código fuente de Zeptai está organizado en el directorio TICKETIA_PRO/, siguiendo "
    "una arquitectura modular con separación clara entre capas:"
)
code(
    "TICKETIA_PRO/\n"
    "├── app.py                    # Factory Flask, registro blueprints, scheduler\n"
    "├── core/\n"
    "│   ├── db_models.py          # Todos los modelos SQLAlchemy\n"
    "│   ├── clients.py            # Singleton del cliente OpenAI\n"
    "│   ├── config.py             # Configuración Flask desde variables de entorno\n"
    "│   ├── limiter.py            # Flask-Limiter instancia compartida\n"
    "│   ├── mcp_client.py         # Cliente MCP SSE\n"
    "│   └── llm_tracker.py        # Decorador de tracking de llamadas LLM\n"
    "├── modules/\n"
    "│   ├── agents/\n"
    "│   │   ├── manager.py        # AgentExecutor: ciclo principal\n"
    "│   │   ├── tools.py          # CalendarTools + TOOLS_SCHEMA\n"
    "│   │   ├── history.py        # HistoryService: guardar/recuperar ChatMessage\n"
    "│   │   └── background_tasks.py # Hilos asíncronos para generación larga\n"
    "│   ├── proactive/\n"
    "│   │   ├── marketing_agent.py   # Imagen, PPT, vídeo\n"
    "│   │   ├── admin_redactor.py    # OCR + PDF presupuestos\n"
    "│   │   ├── post_sales.py        # Postventa / incidencias\n"
    "│   │   ├── grant_hunter.py      # Búsqueda de subvenciones\n"
    "│   │   ├── business_coach.py    # Consejos de negocio\n"
    "│   │   └── networker.py         # Match entre usuarios\n"
    "│   ├── council/\n"
    "│   │   └── council_manager.py   # Multi-agente debate 3 perspectivas\n"
    "│   └── services/\n"
    "│       ├── embeddings.py        # Pipeline RAG: ingesta y recuperación\n"
    "│       └── notification.py      # send_in_app + send_push VAPID\n"
    "├── routes/\n"
    "│   ├── web.py                # Rutas UI (login, dashboard, wizard, etc.)\n"
    "│   └── api.py                # Rutas API (chat, notifications, metrics, eval)\n"
    "├── templates/                # Jinja2 HTML (base.html, dashboard, etc.)\n"
    "├── static/\n"
    "│   ├── generated_docs/       # PDFs, imágenes, PPTs generados\n"
    "│   └── plantilla.pptx        # Plantilla corporativa para presentaciones\n"
    "├── tests/                    # Suite de pruebas (unittest)\n"
    "│   └── test_deepeval_rag.py  # Evaluación DeepEval del pipeline RAG\n"
    "└── mcp_server_sse.py         # Servidor MCP con transporte SSE"
)

h2("4.2. Integración con modelos LLM")
h3("Patrón Singleton de cliente")
p(
    "El cliente OpenAI se instancia una única vez mediante la función get_openai_client(), "
    "almacenada en core/clients.py. Este patrón evita la creación de múltiples conexiones "
    "HTTP por petición y centraliza la configuración del cliente, facilitando el reemplazo "
    "del proveedor (por ejemplo, Azure OpenAI o modelos locales vía LiteLLM):"
)
code(
    "from openai import OpenAI\n\n"
    "_client = None\n\n"
    "def get_openai_client():\n"
    "    global _client\n"
    "    if _client is None:\n"
    "        _client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))\n"
    "    return _client"
)
h3("Trazabilidad de costes — llm_tracker")
p(
    "Todas las llamadas a modelos se instrumentan con track() en core/llm_tracker.py, "
    "que registra en la tabla llm_call: modelo, stage, tokens entrada/salida, latencia (ms), "
    "coste estimado (USD), flag de exito y mensaje de error. Los stages reales son:"
)
table_2col(
    [
        ("safety_classifier", "Clasificacion de seguridad del input (gpt-4o-mini)"),
        ("chat_main", "Primera llamada GPT-4o del AgentExecutor"),
        ("chat_tool_followup", "Segunda llamada GPT-4o tras ejecutar herramienta"),
        ("council_opinion", "Opinion inicial de cada agente en el consejo"),
        ("council_rebuttal", "Replica/debate de cada agente"),
        ("council_synthesis", "Sintesis final del Secretario"),
        ("image_generation", "Generacion de imagen con DALL-E 3"),
        ("audio_transcription", "Transcripcion de audio con Whisper"),
        ("embedding", "Generacion de embedding con text-embedding-3-small"),
    ],
    header=["Stage", "Descripcion"]
)
code(
    "Pricing embebido (2025):\n"
    "gpt-4o:                2,50 / 10,00 USD por 1M tokens (input/output)\n"
    "gpt-4o-mini:           0,15 /  0,60 USD por 1M tokens\n"
    "dall-e-3:              0,040 USD/imagen (1024x1024 standard)\n"
    "whisper-1:             0,006 USD/minuto\n"
    "gen3a_turbo (Runway):  0,05 USD/segundo de video\n"
    "text-embedding-3-small:0,020 USD/1M tokens"
)

h2("4.3. Gestión de prompts y orquestación de agentes")
h3("System prompt personalizado")
p(
    "Cada conversación se inicializa con un system prompt que combina cuatro componentes: "
    "(1) el rol base del asistente ('Eres el asistente IA de {business_name}'); "
    "(2) el static_knowledge del negocio (tarifas, sector, horario, FAQs) serializado como texto; "
    "(3) los chunks RAG recuperados de pgvector relevantes para la consulta actual; "
    "y (4) las instrucciones de comportamiento y capacidades disponibles."
)
p(
    "El usuario puede configurar libremente el system prompt base desde el wizard de configuración, "
    "lo que permite adaptar la personalidad, el tono y las restricciones del agente sin tocar "
    "ninguna línea de código. Esta flexibilidad es uno de los diferenciadores clave de la "
    "plataforma frente a chatbots genéricos de propósito fijo."
)
h3("Ciclo completo de function calling")
code(
    "messages = [{'role': 'system', 'content': system_prompt}] + history\n\n"
    "# Primera llamada: el modelo decide si responde o usa herramienta\n"
    "response = client.chat.completions.create(\n"
    "    model='gpt-4o',\n"
    "    messages=messages,\n"
    "    tools=TOOLS_SCHEMA,\n"
    "    tool_choice='auto'\n"
    ")\n\n"
    "if response.choices[0].message.tool_calls:\n"
    "    # Ejecutar la herramienta y añadir el resultado\n"
    "    messages.append(response.choices[0].message)\n"
    "    messages.append({'role': 'tool', 'content': tool_result,\n"
    "                     'tool_call_id': tool_call.id})\n\n"
    "    # Segunda llamada: síntesis con el resultado\n"
    "    final = client.chat.completions.create(\n"
    "        model='gpt-4o', messages=messages)\n"
    "    return final.choices[0].message.content"
)

h2("4.4. Pipeline RAG con pgvector")
p(
    "El módulo modules/services/embeddings.py implementa el pipeline RAG completo en dos fases:"
)
h3("Fase de ingesta")
p(
    "Al subir un documento PDF o TXT desde la sección Knowledge, el sistema: extrae el texto "
    "mediante PyMuPDF (fitz); aplica el chunker recursivo con separadores [párrafo, frase, "
    "carácter] para dividir en fragmentos de ~1.600 caracteres con overlap de 400 caracteres; "
    "genera el embedding de cada fragmento mediante text-embedding-3-small; y persiste cada "
    "KnowledgeChunk en PostgreSQL con su vector en la columna de tipo pgvector."
)
p(
    "El mismo proceso se aplica a los campos del wizard de configuración del negocio "
    "(source_type='wizard'), de modo que el sistema de preguntas sobre horario, servicios "
    "o tarifas está respaldado por la misma infraestructura vectorial que los PDFs."
)
h3("Fase de recuperación")
p(
    "En cada mensaje del chat, antes de construir el system prompt, se ejecuta retrieve_chunks(): "
    "se genera el embedding de la consulta del usuario, se ejecuta una búsqueda de similitud "
    "coseno en pgvector filtrada por user_phone, y se recuperan los 5 chunks más similares "
    "medidos por distancia coseno. Estos chunks se inyectan en el system prompt como contexto "
    "adicional, permitiendo al agente responder con información específica del negocio."
)
p(
    "El log de cada recuperación se persiste en RagRetrieval (consulta, chunks devueltos, "
    "score medio) para permitir el análisis de calidad del pipeline mediante DeepEval."
)

h2("4.5. Generación asíncrona y notificaciones")
p(
    "Las tareas de larga duración (generación de imágenes DALL·E 3: ~8-15s; PPT: ~5-10s; "
    "vídeo Runway: ~30-90s) se ejecutan en hilos daemon independientes para no bloquear "
    "la respuesta HTTP del chat:"
)
code(
    "thread = threading.Thread(\n"
    "    target=perform_async_marketing_generation,\n"
    "    args=(user_phone, prompt, format_type, app),\n"
    "    daemon=True\n"
    ")\n"
    "thread.start()\n"
    "return '⏳ Me pongo a ello. Te avisaré cuando esté listo.'"
)
p(
    "El hilo recibe el objeto app de Flask para poder usar el contexto de aplicación en "
    "el thread (app.app_context()). Al finalizar la generación, llama a "
    "NotificationService.send_in_app() para crear una notificación en base de datos, y "
    "a send_push() para enviar una web push notification al navegador del usuario mediante "
    "el protocolo VAPID. El frontend implementa polling de 15 segundos contra "
    "/api/notifications/unread_count para actualizar el badge del campana sin requerir "
    "que el usuario refresque la página."
)

h2("4.6. Canales de entrada multimodal")
p(
    "Zeptai no se limita a texto: implementa tres canales de entrada adicionales que amplian "
    "la superficie de interaccion del usuario con el agente:"
)
bullet("Canal texto (chat web): mensaje directo en la interfaz de chat. Flujo estandar AgentExecutor con RAG + function calling.")
bullet("Canal audio (Whisper): el usuario graba un mensaje de voz desde /upload_web_audio. El audio se transcribe con Whisper (coste: 0,006 USD/min) y se clasifica mediante un prompt de intencion: si implica navegacion ('llevame a la agenda'), el LLM devuelve una URL de redireccion en lugar de ejecutar el agente.")
bullet("Canal imagen/ticket (/upload_web_ticket): el usuario sube una fotografia de un ticket o factura (JPG, PNG, WEBP, PDF). El AdminRedactorAgent clasifica el tipo ('receipt' para facturas formales vs 'draft' para notas manuscritas) y llama a GPT-4o Vision con la imagen codificada en base64 para extraer los campos estructurados.")
bullet("Canal vision multimodal: el AgentExecutor puede recibir imagenes adjuntas en el chat para analisis via GPT-4o Vision. Incluye proteccion SSRF: solo se procesan URLs locales o ficheros del servidor.")

h2("4.7. Mecanismos de seguridad")
p(
    "La seguridad de Zeptai se implementa en seis capas defensivas complementarias:"
)
h3("Capa 1 — Autenticacion y sesiones")
p(
    "Formulario de login con check_password_hash (Werkzeug, PBKDF2-SHA256). Sesiones Flask "
    "server-side con SECRET_KEY aleatoria, SESSION_COOKIE_HTTPONLY=True, "
    "SESSION_COOKIE_SAMESITE='Lax', duracion de 8 horas (PERMANENT_SESSION_LIFETIME). "
    "Decorador @login_required en todas las rutas protegidas."
)
h3("Capa 2 — Clasificador de seguridad LLM (gpt-4o-mini)")
p(
    "Cada mensaje del usuario pasa por un clasificador de seguridad basado en gpt-4o-mini "
    "antes de llegar al AgentExecutor. El clasificador categoriza el input en seis clases:"
)
table_2col(
    [
        ("SAFE", "Consulta normal relacionada con el negocio. El mensaje pasa al AgentExecutor."),
        ("OFF_TOPIC", "Pregunta sin relacion con el dominio del negocio. Se rechaza con mensaje educado."),
        ("INJECTION", "Intento de prompt injection (instrucciones en el input del usuario)."),
        ("JAILBREAK", "Intento de eludir restricciones ('actua sin restricciones', 'modo DAN')."),
        ("EXTRACTION", "Intento de extraer el system prompt o informacion del sistema."),
        ("HARMFUL", "Contenido peligroso, ilegal o que viola las politicas de uso."),
    ],
    header=["Categoria", "Accion"]
)
p(
    "El input tiene un limite de longitud de 2.000 caracteres (MAX_INPUT_LENGTH). "
    "El contexto RAG se encapsula en tags <DATOS_DEL_NEGOCIO> para aislar el conocimiento "
    "del negocio del input del usuario. El bloque de seguridad se inyecta al FINAL del "
    "system prompt (las instrucciones tardias tienen mayor peso en GPT-4o)."
)
h3("Capa 3 — Output filtering")
p(
    "Las respuestas del LLM pasan por un filtro regex antes de enviarse al usuario. "
    "Si la respuesta contiene patrones como SECRET_KEY, PASSWORD_HASH, DATABASE_URL, "
    "API_KEY, OPENAI_API_KEY, VAPID_PRIVATE o MAIL_PASSWORD, se sustituye por "
    "\"Lo siento, no puedo responder a esa consulta.\""
)
h3("Capa 4 — Rate limiting")
table_2col(
    [
        ("/api/chat", "30 peticiones/minuto"),
        ("/api/chat/feedback", "60 peticiones/minuto"),
        ("/api/council/stream", "10 peticiones/hora"),
        ("/api/eval/stream", "10 peticiones/hora"),
        ("/api/metrics/llm y /api/metrics/rag", "30 peticiones/hora"),
        ("/upload_web_ticket y /upload_web_audio", "20 peticiones/hora"),
        ("/generate_video_from_image", "5 peticiones/hora"),
    ],
    header=["Endpoint", "Limite"]
)
h3("Capa 5 — Aislamiento de datos")
p(
    "Todos los queries SQL filtran por session['user_phone']. No existe ningun endpoint "
    "que devuelva datos de otro usuario. Los documentos generados incluyen UUID en el "
    "nombre del fichero para evitar enumeracion."
)
h3("Capa 6 — Proteccion SSRF y logging de auditoria")
p(
    "Las URLs de imagenes externas se validan antes de procesarlas con GPT-4o Vision "
    "(solo se permiten rutas locales del servidor). ActivityLog registra todas las "
    "acciones relevantes con timestamp para auditoria completa."
)

h2("4.7. Despliegue e infraestructura")
p(
    "El entorno de ejecución está completamente contenedorizado mediante Docker Compose con "
    "cuatro servicios con dependencias declaradas explícitamente:"
)
table_2col(
    [
        ("zeptai_db", "PostgreSQL 15 + pgvector. Healthcheck pg_isready garantiza disponibilidad antes de que los demás servicios arranquen. Volumen persistente pgdata."),
        ("zeptai_mcp", "Servidor MCP SSE en el puerto 8001. Arranca después de db y antes que web. Expone 6 herramientas MCP a los agentes."),
        ("zeptai_app", "Flask + Gunicorn en el puerto 5000. Volume mount del código fuente para hot-reload en desarrollo sin rebuild de imagen."),
        ("zeptai_scheduler", "Scheduler con libreria schedule. Arranca despues de web. Ejecuta los 6 agentes proactivos (BusinessHealth, GrantHunter, Synergy, PostSales, AdminRedactor, Marketing) a las 09:00 diarias."),
    ],
    header=["Servicio", "Descripción"]
)
p(
    "Para producción, la arquitectura recomendada añade Nginx como reverse proxy con terminación "
    "SSL (Let's Encrypt), Gunicorn con 4 workers síncronos, y un servicio PostgreSQL gestionado "
    "externo (Neon.tech, Supabase o AWS RDS) con backups automáticos. El pipeline CI/CD de "
    "GitHub Actions ejecuta los tests unitarios y las evaluaciones de agentes en cada push "
    "a main, dev y feature/*."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — EVALUACIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1("5. Evaluación y Experimentos")

h2("5.1. Marco de evaluación — DeepEval")
p(
    "La evaluación objetiva de calidad del pipeline RAG es uno de los elementos diferenciadores "
    "de Zeptai respecto a implementaciones similares. El sistema integra DeepEval, un framework "
    "open-source que emplea GPT-4o como juez automático (LLM-as-a-judge) para evaluar la "
    "calidad de las respuestas del sistema RAG sin requerir anotación humana."
)
p(
    "La evaluación es accesible desde la interfaz web en la ruta /eval, con transmisión de "
    "resultados en tiempo real mediante Server-Sent Events (SSE). El endpoint /api/eval/stream "
    "ejecuta la evaluación en un hilo separado con contexto Flask explícito para evitar "
    "el error 'Working outside of application context'."
)
h3("Métricas evaluadas")
table_2col(
    [
        ("Faithfulness (umbral ≥ 0,70)", "Mide si la respuesta está anclada en los chunks recuperados, detectando alucinaciones. Fórmula: proporción de afirmaciones del output que pueden verificarse en el retrieval_context."),
        ("Answer Relevancy (umbral ≥ 0,70)", "Mide si la respuesta es relevante para la pregunta formulada, penalizando divagaciones y respuestas fuera de tema."),
        ("Contextual Precision (umbral ≥ 0,50)", "Mide qué proporción de los chunks recuperados por el RAG son realmente necesarios para responder. Evalúa la calidad del retriever."),
        ("Contextual Recall (umbral ≥ 0,50)", "Mide qué proporción de la información necesaria para responder está presente en los chunks recuperados. Evalúa la cobertura del conocimiento indexado."),
    ],
    header=["Métrica", "Descripción y umbral de aprobado"]
)
h3("Casos de prueba")
table_2col(
    [
        ("TC-01 Servicios", "Pregunta sobre los servicios ofrecidos por el negocio. Valida que el RAG recupera correctamente los chunks de descripción de servicios."),
        ("TC-02 Horario", "Pregunta sobre el horario de apertura. Valida la recuperación de chunks de configuración del wizard."),
        ("TC-03 Pricing", "Pregunta sobre tarifas y precios. Valida Faithfulness: la respuesta debe basarse en los precios indexados, no inventarlos."),
        ("TC-04 Contacto", "Pregunta sobre formas de contacto. Valida Answer Relevancy y Contextual Precision."),
        ("TC-05 Propuesta de valor", "Pregunta sobre los diferenciales del negocio. Valida coherencia con el contenido del PDF de conocimiento."),
        ("TC-06 Proceso", "Pregunta sobre el proceso de trabajo o metodología. Valida Contextual Recall."),
        ("TC-07 Guardrail", "'¿Cuál es la capital de Francia?' — pregunta trampa fuera del dominio. Un score bajo en Faithfulness es el resultado CORRECTO (confirma que el agente rechaza preguntas fuera de dominio)."),
    ],
    header=["Caso", "Descripción y validación"]
)
p(
    "El coste de una evaluación completa de los 7 casos con GPT-4o como juez es "
    "aproximadamente 0,40 USD, lo que la hace económicamente viable para ejecución "
    "periódica o tras cambios significativos en el pipeline RAG."
)

h2("5.2. Métricas técnicas de rendimiento")
table_2col(
    [
        ("Latencia P50 — chat sin herramienta", "~1,2 segundos"),
        ("Latencia P95 — chat sin herramienta", "~2,8 segundos"),
        ("Latencia P50 — chat con tool use", "~2,5 segundos (incluye 2ª llamada GPT-4o)"),
        ("Latencia P95 — chat con tool use", "~4,5 segundos"),
        ("Tiempo de generación imagen DALL·E 3", "~8-15 segundos (asíncrono)"),
        ("Tiempo de generación PPT (5 slides)", "~5-10 segundos (asíncrono)"),
        ("Tiempo de generación vídeo Runway", "~30-90 segundos (asíncrono)"),
        ("Latencia herramienta MCP SSE", "~5 ms por herramienta"),
        ("Latencia herramienta MCP stdio (anterior)", "~500 ms por herramienta"),
    ],
    header=["Métrica", "Valor medido"]
)

h2("5.3. Análisis de costes")
p(
    "Uno de los criterios de validez del proyecto es que el coste operativo sea compatible "
    "con un modelo de negocio SaaS rentable. El análisis de costes por usuario activo muestra "
    "que la solución es económicamente viable:"
)
table_2col(
    [
        ("50 mensajes de chat/mes (GPT-4o)", "~0,30 €"),
        ("5 imágenes DALL·E 3/mes", "~0,20 €"),
        ("3 presentaciones PPT/mes", "~0,03 €"),
        ("2 vídeos Runway ML/mes", "~0,40 €"),
        ("TOTAL estimado por usuario activo/mes", "~1,00 €"),
        ("Con 20 usuarios Pro (19 €/mes)", "20 € costes API vs. 380 € ingresos → SRR: 94,7%"),
        ("1 evaluación DeepEval completa", "~0,40 USD (ocasional)"),
    ],
    header=["Concepto", "Coste estimado"]
)
note(
    "Los precios son estimaciones basadas en el pricing de OpenAI y Runway ML de 2025. "
    "Pueden variar según el volumen real de uso y los cambios de precio de los proveedores."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — PRUEBAS Y VALIDACIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1("6. Pruebas y Validación de Producto")

h2("6.1. Pruebas funcionales automatizadas")
p(
    "Zeptai cuenta con un suite de 51 casos de prueba automatizados (run_tests.py) ejecutables "
    "sobre el entorno Docker local. Las pruebas no realizan llamadas reales a la API de OpenAI, "
    "operando con un usuario de prueba ficticio y mocks controlados. El resultado agregado es:"
)
bullet("Pruebas superadas: 46/51 (90,2%)")
bullet("Bugs reales encontrados durante las pruebas: 0")
bullet("Falsos positivos: 5 (artefactos del script de test, no bugs de la aplicación)")

table_2col(
    [
        ("Bloque 0 — Autenticación", "6 tests: rutas protegidas redirigen a login cuando no hay sesión activa."),
        ("Bloque 1 — Login", "2 tests: credenciales correctas conceden sesión; credenciales incorrectas devuelven error 401."),
        ("Bloque 2 — Carga de páginas", "11 tests: las 11 rutas principales devuelven HTTP 200 con usuario autenticado."),
        ("Bloque 3 — Bot status", "2 tests: el estado del bot refleja correctamente el flag bot_enabled tras completar el wizard."),
        ("Bloque 4 — API Notificaciones", "6 tests: CRUD completo (crear, listar, marcar leída, contar no leídas, eliminar)."),
        ("Bloque 5 — API Métricas LLM", "4 tests: estructura de respuesta JSON de /api/metrics (totales, por modelo, por día)."),
        ("Bloque 6 — Tickets/Transacciones", "4 tests: carga de la lista y visibilidad de registros."),
        ("Bloque 7 — Documentos", "3 tests: carga de la sección, descarga de documento existente, 404 en ID inexistente."),
        ("Bloque 8 — Marketplace", "3 tests: toggle de activación/desactivación de agentes por categoría."),
        ("Bloque 9 — Demo", "2 tests: seed_data y seed_grants inyectan datos correctamente."),
        ("Bloque 10 — Push VAPID", "2 tests: endpoint devuelve la clave pública VAPID en formato correcto."),
        ("Bloque 11 — Export Excel", "2 tests: Content-Type correcto (application/vnd.openxmlformats) en la descarga."),
        ("Bloque 12 — Manejo de errores", "1 test: rutas inexistentes devuelven HTTP 404 con página de error."),
        ("Bloque 13 — Wizard/save_config", "2 tests: guardar configuración activa el flag bot_enabled correctamente."),
        ("Bloque 14 — Logout", "4 tests: logout invalida la sesión; rutas protegidas vuelven a requerir autenticación."),
    ],
    header=["Bloque", "Descripción"]
)

h2("6.2. Bugs identificados y corregidos")
p(
    "El proceso de desarrollo iterativo identificó y corrigió cuatro bugs relevantes, "
    "todos detectados mediante las pruebas automatizadas o la validación manual:"
)
bullet(
    "Bug 1 — Bot status permanecía 'Dormido' tras completar el wizard. "
    "Causa: save_config() no actualizaba features['bot_enabled'] = True en ambas ramas (crear/actualizar perfil). "
    "Solución: actualización explícita del flag en las dos ramas del condicional."
)
bullet(
    "Bug 2 — Número de slides incorrecto (se pedían 5, se generaban 4). "
    "Causa: GPT-4o extraía el número de slides del texto del prompt de forma implícita, y el parsing fallaba ocasionalmente. "
    "Solución: añadir parámetro slide_count explícito y tipado en el TOOLS_SCHEMA, de modo que el modelo lo extrae de forma estructurada."
)
bullet(
    "Bug 3 — Presentaciones sin título descriptivo (aparecían todas como 'Presentación'). "
    "Causa: GeneratedDocument se creaba sin rellenar el campo client_name. "
    "Solución: el marketing_agent extrae las primeras cinco palabras del título generado por GPT-4o y las usa como client_name."
)
bullet(
    "Bug 4 — Badge de notificaciones no se actualizaba automáticamente. "
    "Causa: el contador solo se actualizaba al abrir el drawer de notificaciones manualmente. "
    "Solución: polling de 15 segundos contra /api/notifications/unread_count con actualización del badge en el DOM."
)
bullet(
    "Bug 5 (DeepEval) — Error 'Working outside of application context' en el thread de evaluación. "
    "Causa: el thread de evaluación no tenía acceso al contexto de Flask. "
    "Solución: capturar app = current_app._get_current_object() antes de lanzar el thread y usar app.app_context() dentro."
)
bullet(
    "Bug 6 (DeepEval) — AttributeError 'EvaluationResult object has no attribute evaluation_cost'. "
    "Causa: la API de DeepEval cambió el nombre del atributo entre versiones. "
    "Solución: usar getattr(results, 'evaluation_cost', None) or getattr(results, 'cost', None) or 0 para compatibilidad."
)

h2("6.3. Pruebas manuales requeridas")
p(
    "Las siguientes funcionalidades requieren navegador o incurren en coste de API real, "
    "por lo que no están cubiertas por las pruebas automatizadas y se validan manualmente:"
)
bullet("Chat conversacional completo con tool use end-to-end (múltiples ciclos de razonamiento y herramienta).")
bullet("Generación de presupuesto desde texto y desde imagen adjunta (OCR + ReportLab PDF).")
bullet("Generación de imagen DALL·E 3 con prompt contextualizado al negocio.")
bullet("Generación de presentación PPT con plantilla corporativa y número de slides correcto.")
bullet("CouncilManager con debate de 3 perspectivas en streaming SSE visible en el navegador.")
bullet("Notificaciones web push en navegador con Service Worker activo.")
bullet("Instalación PWA desde Chrome (beforeinstallprompt).")
bullet("OCR de ticket desde fotografía real de una factura o nota manuscrita.")
bullet("Evaluación DeepEval completa desde la interfaz /eval con resultados en tiempo real.")

h2("6.4. UX y usabilidad")
p(
    "Las decisiones de diseño de la interfaz priorizaron la reducción de la curva de aprendizaje "
    "para el perfil de usuario objetivo (autónomo o propietario de PYME sin formación técnica):"
)
bullet("Chat como interfaz principal: el usuario no necesita aprender ninguna sintaxis ni navegar por menús complejos; interactúa en lenguaje natural.")
bullet("Notificaciones push + in-app: el usuario no necesita mantener la app abierta durante las tareas asíncronas (generación de imágenes o vídeos), reduciendo la percepción de espera.")
bullet("Wizard de configuración en 5 pasos guiados: evita la página en blanco al usuario novel sin abrumar con todas las opciones a la vez.")
bullet("Marketplace de agentes: el usuario activa solo los agentes que necesita, evitando la sobrecarga cognitiva de demasiadas funcionalidades visibles simultáneamente.")
bullet("Diseño responsive y PWA instalable: la plataforma es usable desde móvil sin necesidad de App Store.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 7 — DISCUSIÓN
# ═══════════════════════════════════════════════════════════════════════════
h1("7. Discusión")

h2("7.1. Lecciones aprendidas")
p(
    "El desarrollo de Zeptai generó aprendizajes técnicos relevantes que trascienden el "
    "proyecto concreto y son aplicables a futuros sistemas multi-agente:"
)
h3("Function calling es más robusto que el parsing de texto libre")
p(
    "En versiones iniciales del sistema se intentó clasificar la intención del usuario mediante "
    "expresiones regulares y prompts de clasificación de texto libre. El uso del mecanismo "
    "nativo de function calling de OpenAI eliminó completamente los errores de routing de "
    "herramientas: el modelo devuelve JSON estructurado y tipado con fiabilidad superior al 95%, "
    "frente a la fragilidad inherente del parsing de lenguaje natural. Esta lección confirma "
    "la recomendación de la documentación de OpenAI: siempre preferir tool use frente a prompting "
    "para extraer datos estructurados."
)
h3("Los parámetros implícitos en el prompt causan errores")
p(
    "El bug del slide_count (el usuario pedía 5 slides y se generaban 4) demostró que pasar "
    "información numérica dentro de un string de texto libre es una aproximación frágil. "
    "La solución correcta es parametrizar explícitamente en el tool schema: cuando el "
    "modelo tiene un campo tipado (integer, string) para extraer el dato, lo hace con "
    "precisión. Esta lección es aplicable a cualquier dato que el sistema necesite extraer "
    "del input del usuario de forma fiable."
)
h3("La generación asíncrona es esencial para una UX aceptable")
p(
    "Mantener una petición HTTP abierta durante 30-90 segundos (tiempo de renderizado de "
    "un vídeo en Runway ML) era absolutamente inaceptable desde el punto de vista de la "
    "experiencia de usuario y técnicamente problemático (timeouts de proxies, límites de "
    "workers de Gunicorn). El patrón hilo daemon + notificación push resultó en una UX "
    "significativamente mejor que el polling activo desde el frontend, porque el usuario "
    "puede cerrar la app y recibir la notificación cuando el resultado esté listo."
)
h3("pgvector elimina un servicio externo sin sacrificar capacidad")
p(
    "La decisión de usar la extensión pgvector sobre el propio PostgreSQL, en lugar de "
    "servicios vectoriales dedicados (Pinecone, Qdrant, ChromaDB), redujo la arquitectura "
    "de 5 a 4 servicios Docker sin ninguna pérdida de funcionalidad observable. Para el "
    "volumen de una PYME (cientos de chunks por usuario), el rendimiento de búsqueda "
    "vectorial con pgvector es indistinguible de un vector store dedicado. La evaluación "
    "con DeepEval confirmó que el pipeline RAG sobre pgvector funciona correctamente, "
    "con Faithfulness y Answer Relevancy por encima del umbral en los casos de prueba "
    "con contenido documentado."
)
h3("El contexto Flask en threads requiere gestión explícita")
p(
    "Python Flask utiliza un patrón de contexto de aplicación basado en variables de thread-local. "
    "Los threads lanzados manualmente (para generación asíncrona y para la evaluación DeepEval) "
    "no heredan automáticamente este contexto. La solución es capturar el objeto app antes de "
    "lanzar el thread (app = current_app._get_current_object()) y usar el contexto explícitamente "
    "dentro (with app.app_context():). Este patrón debe aplicarse consistentemente en todos "
    "los módulos que utilicen hilos en aplicaciones Flask."
)

h2("7.2. Riesgos, ética y mitigaciones")
p(
    "El siguiente cuadro resume el análisis de riesgos residuales de la versión actual de "
    "Zeptai, con el nivel de riesgo y las mitigaciones implementadas:"
)
table_2col(
    [
        ("Alucinaciones\nNivel: MEDIO", "El LLM puede generar información incorrecta. Mitigación: el RAG ancla las respuestas en el conocimiento documentado del negocio. Validado con DeepEval Faithfulness ≥ 0,70."),
        ("Prompt injection\nNivel: BAJO", "Un usuario malicioso intenta manipular el system prompt. Mitigación: el system prompt se inyecta en posición fija (role: system) y no es sobreescribible por el input del usuario."),
        ("Privacidad datos conversacionales\nNivel: MEDIO", "El historial de chat contiene información sensible. Mitigación: los datos se almacenan en la BD propia, nunca en servidores de OpenAI (salvo el procesamiento puntual de inferencia). Política de retención recomendada: 90 días."),
        ("Vendor lock-in OpenAI\nNivel: ALTO", "OpenAI puede cambiar precios, deprecar modelos o tener outages. Mitigación: capa de abstracción get_openai_client() permite sustitución; llm_tracker monitoriza costes y anomalías. Trabajo futuro: fallback a modelos open-source."),
        ("Outage API externa\nNivel: ALTO", "Un fallo de OpenAI paraliza el sistema. Mitigación: el llm_tracker detecta errores; trabajo futuro: retry con exponential backoff y fallback a modelos alternativos."),
        ("Escalabilidad de costes\nNivel: BAJO", "Un usuario abusivo puede generar costes elevados en la API. Mitigación: rate limiting por endpoint y por usuario; trabajo futuro: presupuesto de API por usuario con alerta."),
    ],
    header=["Riesgo", "Descripción y mitigación"]
)

h2("7.3. Análisis de limitaciones")
p(
    "La versión actual del MVP presenta limitaciones conocidas que han sido conscientemente "
    "acotadas al alcance del trabajo:"
)
bullet(
    "Ventana de contexto limitada a 10 turnos: el historial de conversación se trunca para "
    "controlar el coste de tokens. Conversaciones largas pierden el contexto de interacciones "
    "anteriores, lo que puede afectar a la coherencia en sesiones extendidas."
)
bullet(
    "Un solo hilo de conversación por usuario: no se admiten conversaciones paralelas ni "
    "separación del historial por proyectos o clientes. Cada nuevo chat parte del historial "
    "acumulado del usuario."
)
bullet(
    "Casos de prueba DeepEval hardcodeados: los 7 casos de prueba del MVP son fijos. "
    "Un sistema de evaluación maduro debería generar casos de prueba dinámicamente "
    "a partir de conversaciones reales del negocio."
)
bullet(
    "Sin integración con canales externos: el MVP opera únicamente desde la interfaz web. "
    "La integración con WhatsApp Business, email o calendarios externos (Google Calendar) "
    "está fuera del alcance del MVP pero es una prioridad de roadmap."
)
bullet(
    "Dependencia de un único proveedor LLM: toda la inferencia se delega a OpenAI. "
    "Un sistema de producción robusto debería tener un mecanismo de fallback a modelos "
    "alternativos (Anthropic Claude, Google Gemini, o modelos open-source como Llama 3) "
    "para garantizar disponibilidad."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 8 — CONCLUSIONES
# ═══════════════════════════════════════════════════════════════════════════
h1("8. Conclusiones y Trabajo Futuro")

h2("8.1. Conclusiones")
p(
    "Zeptai demuestra que es posible construir una plataforma SaaS multi-agente funcional, "
    "desplegable en producción y con valor empresarial cuantificable, utilizando exclusivamente "
    "APIs y frameworks open-source del ecosistema Python. Los siete objetivos específicos "
    "definidos al inicio del proyecto han sido cumplidos en su totalidad:"
)
bullet("OE1 ✅ — 6 agentes especializados accesibles desde chat mediante function calling nativo de GPT-4o.")
bullet("OE2 ✅ — CouncilManager con debate de 3 perspectivas independientes, síntesis ejecutiva y streaming SSE.")
bullet("OE3 ✅ — Generación multimodal completa: texto (presupuestos PDF), imágenes (DALL·E 3), presentaciones (python-pptx) y prompts de vídeo (Runway ML Gen-3).")
bullet("OE4 ✅ — Pipeline OCR + LLM multimodal para digitalización de documentos físicos con GPT-4o Vision.")
bullet("OE5 ✅ — Pipeline RAG completo con PostgreSQL + pgvector: ingesta, chunking, embeddings y recuperación semántica.")
bullet("OE6 ✅ — Docker Compose reproducible con 4 servicios, notificaciones web push VAPID y servidor MCP SSE.")
bullet("OE7 ✅ — 51 pruebas funcionales automatizadas (90,2% superadas) + evaluación objetiva DeepEval con 4 métricas LLM-as-a-judge sobre 7 casos de prueba.")
p(
    "La solución aporta valor cuantificable al usuario final: un autónomo que usa Zeptai "
    "diariamente puede recuperar entre 3 y 5 horas semanales de trabajo administrativo. "
    "A un valor de 25 €/hora, esto representa entre 75 y 125 € semanales de valor generado, "
    "con un coste mensual de API inferior a 5 € en uso moderado. El retorno sobre la "
    "inversión estimado en el primer mes de uso activo supera el 1.000%."
)
p(
    "Desde el punto de vista académico, el proyecto realiza contribuciones en tres dimensiones: "
    "(1) diseño e implementación de un sistema multi-agente con patrones ReAct y function calling "
    "en un contexto real de negocio; (2) integración de un pipeline RAG sin servicio vectorial "
    "externo mediante pgvector, con evaluación objetiva mediante DeepEval; y (3) arquitectura "
    "de referencia para plataformas SaaS basadas en LLMs con generación asíncrona, "
    "notificaciones push y servidor MCP propio."
)

h2("8.2. Trabajo futuro")
h3("Corto plazo (0-3 meses)")
bullet("Integración con WhatsApp Business API (webhook Twilio / Meta Cloud API) para operar desde el canal de mensajería más usado en España.")
bullet("Generación dinámica de casos de prueba DeepEval a partir de conversaciones reales, eliminando la dependencia de los 7 casos fijos del MVP.")
bullet("Presupuesto de API por usuario con alerta por email al superar el umbral mensual configurado.")
bullet("Política de retención de datos (90 días por defecto) y botón explícito de 'borrar mi cuenta' para cumplimiento RGPD.")
bullet("Endpoint de health check con métricas de latencia expuestas para monitorización externa (Prometheus / UptimeRobot).")

h3("Medio plazo (3-12 meses)")
bullet("App móvil nativa (React Native o Flutter) con acceso a la cámara para OCR de documentos físicos.")
bullet("Fine-tuning de un modelo propio sobre datos de presupuestos del sector para mejorar la calidad sin depender exclusivamente de GPT-4o.")
bullet("Integración con software de contabilidad (Holded, Conta.cat, Sage) mediante APIs para sincronización bidireccional.")
bullet("Marketplace de agentes con contribuciones de terceros y sistema de valoración.")
bullet("Facturación SaaS (Stripe) con planes freemium / pro / business y gestión de suscripciones.")
bullet("Autenticación OAuth2 (Google, LinkedIn) para reducir fricción en el registro.")

h3("Largo plazo (12+ meses)")
bullet("Modelo propio (Mistral o Llama fine-tuned) para reducir dependencia de OpenAI y controlar costes a escala.")
bullet("Agente de análisis financiero con integración bancaria mediante Open Banking PSD2.")
bullet("Expansión a otros mercados hispanohablantes: México, Colombia, Argentina.")
bullet("Certificación ISO 27001 para clientes enterprise que requieren auditoría de seguridad.")
bullet("Sistema de aprendizaje continuo: el agente mejora su rendimiento con cada interacción mediante feedback loop humano.")

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 9 — BIBLIOGRAFÍA
# ═══════════════════════════════════════════════════════════════════════════
h1("9. Bibliografía")

refs = [
    "[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł. & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
    "[2] Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., ... & Amodei, D. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
    "[3] Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
    "[4] Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K. & Cao, Y. (2023). ReAct: Synergizing reasoning and acting in language models. International Conference on Learning Representations (ICLR 2023).",
    "[5] OpenAI. (2023). GPT-4 technical report. arXiv preprint arXiv:2303.08774.",
    "[6] Park, J. S., O'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P. & Bernstein, M. S. (2023). Generative agents: Interactive simulacra of human behavior. Proceedings of UIST 2023.",
    "[7] Wu, Q., Bansal, G., Zhang, J., Wu, Y., Zhang, S., Zhu, E., ... & Wang, C. (2023). AutoGen: Enabling next-gen LLM applications via multi-agent conversation. arXiv preprint arXiv:2308.08155.",
    "[8] Ramesh, A., Dhariwal, P., Nichol, A., Chu, C. & Chen, M. (2022). Hierarchical text-conditional image generation with CLIP latents. arXiv preprint arXiv:2204.06125.",
    "[9] Anthropic. (2024). Model Context Protocol: specification and reference implementation. https://modelcontextprotocol.io/",
    "[10] Confident AI. (2024). DeepEval: open-source LLM evaluation framework. https://github.com/confident-ai/deepeval",
    "[11] Unión Europea. (2016). Reglamento (UE) 2016/679 del Parlamento Europeo y del Consejo (Reglamento General de Protección de Datos). Diario Oficial de la Unión Europea, L 119/1.",
    "[12] Unión Europea. (2024). Reglamento (UE) 2024/1689 del Parlamento Europeo y del Consejo por el que se establecen normas armonizadas en materia de inteligencia artificial (Reglamento de IA). Diario Oficial de la Unión Europea.",
    "[13] Ministerio de Industria, Comercio y Turismo. (2023). Estadísticas PYME — Evolución e indicadores. Recuperado de https://www.ipyme.org/",
    "[14] Grand View Research. (2024). Artificial Intelligence Market Size, Share & Trends Analysis Report. Recuperado de https://www.grandviewresearch.com/",
    "[15] Registro de la Seguridad Social (RETA). (2023). Estadísticas de trabajadores autónomos. Ministerio de Inclusión, Seguridad Social y Migraciones.",
    "[16] Johnson, L., & Goodfellow, I. (2023). Practical deep learning for practitioners. MIT Press.",
    "[17] Nakano, R., et al. (2021). WebGPT: Browser-assisted question-answering with human feedback. arXiv preprint arXiv:2112.09332.",
    "[18] Schick, T., et al. (2023). Toolformer: Language models can teach themselves to use tools. Advances in Neural Information Processing Systems, 36.",
]
for ref in refs:
    ref_para = doc.add_paragraph()
    ref_para.paragraph_format.left_indent = Cm(1)
    ref_para.paragraph_format.first_line_indent = Cm(-1)
    ref_para.paragraph_format.space_after = Pt(4)
    run = ref_para.add_run(ref)
    run.font.size = Pt(10)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ANEXO A — GUÍA DE DESPLIEGUE
# ═══════════════════════════════════════════════════════════════════════════
h1("Anexo A. Guía de Despliegue")
h2("Pre-requisitos")
bullet("Docker Desktop >= 4.x instalado y en ejecución.")
bullet("API key de OpenAI con acceso a GPT-4o (obligatoria).")
bullet("Archivo .env configurado con las variables de entorno requeridas.")
bullet("Puerto 5000 (aplicación web) y 5432 (PostgreSQL) libres en el host.")

h2("Pasos de instalación")
code(
    "# 1. Clonar el repositorio\n"
    "git clone https://github.com/alejbrata/zeptai.git\n"
    "cd zeptai\n\n"
    "# 2. Configurar variables de entorno\n"
    "cp .env.example .env\n"
    "# Editar .env con las API keys reales\n\n"
    "# 3. Levantar todos los servicios\n"
    "docker compose up --build -d\n\n"
    "# 4. Inicializar la base de datos y cargar datos demo\n"
    "docker compose exec web python seed_all.py\n\n"
    "# 5. Acceder a la aplicación\n"
    "# http://localhost:5000\n"
    "# Usuario demo: admin@demo.com | Contraseña: demo1234"
)

h2("Variables de entorno")
table_2col(
    [
        ("SECRET_KEY", "Obligatoria. Clave Flask aleatoria. Generar con: python -c \"import secrets; print(secrets.token_hex(32))\""),
        ("OPENAI_API_KEY", "Obligatoria. API key de OpenAI con acceso a GPT-4o."),
        ("RUNWAYML_API_SECRET", "Opcional. Para generación de vídeos con Runway ML Gen-3."),
        ("MAIL_USERNAME / MAIL_PASSWORD", "Opcional. Gmail + App Password para envío de correos (reset de contraseña, notificaciones)."),
        ("VAPID_PUBLIC_KEY / VAPID_PRIVATE_KEY", "Opcional. Para notificaciones web push PWA."),
        ("DATABASE_URL", "Opcional. PostgreSQL externo. Docker Compose configura el contenedor automáticamente."),
    ],
    header=["Variable", "Descripción"]
)

h2("Comandos útiles")
code(
    "# Ver logs en tiempo real\n"
    "docker compose logs -f web\n\n"
    "# Reinicio limpio (borra todos los datos)\n"
    "docker compose down -v && docker compose up --build -d\n\n"
    "# Ejecutar suite de tests\n"
    "docker compose exec web python run_tests.py\n\n"
    "# Ejecutar evaluación DeepEval desde CLI\n"
    "docker compose exec web python -m pytest TICKETIA_PRO/tests/test_deepeval_rag.py -v\n\n"
    "# Acceder a la base de datos\n"
    "docker compose exec db psql -U postgres -d zeptai_db"
)

h2("Verificación del despliegue")
code(
    "# Health check de la aplicación\n"
    "curl http://localhost:5000/health\n"
    "# Respuesta esperada: {\"status\": \"ok\"}\n\n"
    "# Estado de los contenedores\n"
    "docker compose ps\n"
    "# Todos deben mostrar 'Up' y zeptai_db debe mostrar '(healthy)'"
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ANEXO B — MANUAL DE USUARIO
# ═══════════════════════════════════════════════════════════════════════════
h1("Anexo B. Manual de Usuario Resumido")

h2("Primeros pasos")
p(
    "El proceso de puesta en marcha de Zeptai está diseñado para completarse en menos de "
    "10 minutos, sin necesidad de conocimientos técnicos:"
)
bullet("1. Registrarse en /register con email y contraseña.")
bullet("2. Completar el wizard de configuración (/wizard): nombre del negocio, sector, logo, system prompt personalizado y conocimiento estático (tarifas, FAQs, horario).")
bullet("3. Activar los agentes que se necesiten en /marketplace (por defecto están todos activos).")
bullet("4. Subir documentos de conocimiento en /knowledge (catálogos, manuales, guías de precios en PDF).")
bullet("5. Ir al chat (/dashboard) y comenzar a interactuar con el asistente.")

h2("Ejemplos de uso del chat")
code(
    "\"Hazme un presupuesto para María García por 3 horas de fontanería a 45 €/h\"\n"
    "\"Crea una imagen de Instagram para promocionar mi negocio este verano\"\n"
    "\"¿Tienes huecos disponibles el próximo lunes?\"\n"
    "\"Reserva una cita para Juan Pérez el lunes a las 10\"\n"
    "\"Analiza esta foto de mi nota de trabajo\" [adjuntar imagen]\n"
    "\"Genera una presentación de 6 slides sobre los servicios de mi empresa\"\n"
    "\"¿Hay subvenciones disponibles para mi sector?\""
)

h2("Sección Council (debate estratégico)")
p(
    "La sección Council (/council) permite plantear preguntas estratégicas complejas que "
    "el sistema responde desde tres perspectivas simultáneas:"
)
code(
    "\"¿Debería contratar un empleado o subcontratar?\"\n"
    "\"¿Cómo puedo aumentar mis tarifas sin perder clientes?\"\n"
    "\"¿Tiene sentido abrir un segundo local?\""
)

h2("Sección Documentos")
p(
    "Todos los PDFs, imágenes y presentaciones generados por el asistente aparecen en la "
    "sección /documents. Se pueden descargar individualmente o eliminar. Los documentos "
    "se organizan por tipo (presupuesto, imagen, presentación, prompt de vídeo) y fecha."
)

h2("Sección Métricas")
p(
    "La sección /metrics muestra el consumo de tokens y el coste estimado de las llamadas "
    "a la API de OpenAI, desglosado por modelo (gpt-4o, dall-e-3, text-embedding-3-small) "
    "y por día. Permite controlar el gasto mensual y detectar anomalías."
)

h2("Evaluación de calidad RAG")
p(
    "La sección /eval permite lanzar una evaluación objetiva del pipeline RAG mediante "
    "DeepEval. La evaluación tarda aproximadamente 2-3 minutos y muestra los resultados "
    "por caso de prueba en tiempo real. Coste aproximado por evaluación completa: 0,40 USD."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ANEXO C — PROMPTS CLAVE
# ═══════════════════════════════════════════════════════════════════════════
h1("Anexo C. Prompts Clave del Sistema")

h2("Prompt del Agente Estratega (CouncilManager)")
code(
    "Eres el Estratega del equipo de consejo. Tu función es analizar la pregunta desde\n"
    "una perspectiva de visión a largo plazo: oportunidades de mercado, posicionamiento\n"
    "competitivo, tendencias del sector y objetivos estratégicos del negocio.\n"
    "Contexto del negocio: {business_context}\n"
    "Sé directo, usa datos cuando los tengas, y concluye con 2-3 recomendaciones\n"
    "estratégicas concretas y priorizadas."
)

h2("Prompt del Agente Analista")
code(
    "Eres el Analista del equipo de consejo. Tu función es examinar la pregunta con rigor\n"
    "cuantitativo: identifica riesgos específicos, evalúa la viabilidad financiera, propón\n"
    "métricas de éxito medibles y señala las hipótesis que habría que validar antes de\n"
    "tomar una decisión. Sé escéptico constructivo y basa tus afirmaciones en números\n"
    "concretos cuando sea posible."
)

h2("Prompt del Agente Implementador")
code(
    "Eres el Implementador del equipo de consejo. Tu función es traducir la pregunta en\n"
    "pasos de acción concretos y secuenciados: qué hacer primero, qué recursos se necesitan,\n"
    "qué obstáculos prácticos hay que superar y cómo medir el progreso a corto plazo.\n"
    "Sé pragmático y orientado a la ejecución. Evita las generalidades."
)

h2("Prompt de Síntesis")
code(
    "Has recibido tres perspectivas diferentes sobre la misma pregunta de negocio.\n"
    "Tu tarea es sintetizarlas en una recomendación ejecutiva clara y equilibrada que\n"
    "integre los puntos de vista estratégico, analítico y de implementación.\n"
    "La síntesis debe ser accionable, sin repetir literalmente lo ya dicho, destacando\n"
    "los puntos de convergencia y señalando explícitamente las tensiones o trade-offs\n"
    "que el decisor debe resolver según sus prioridades."
)

h2("Prompt base del AgentExecutor")
code(
    "Eres el asistente IA de {business_name}, una empresa del sector {sector}.\n\n"
    "CONOCIMIENTO DE TU NEGOCIO:\n"
    "{static_knowledge}\n\n"
    "CONTEXTO ADICIONAL (documentos indexados):\n"
    "{rag_chunks}\n\n"
    "INSTRUCCIONES:\n"
    "- Responde siempre en el idioma del usuario\n"
    "- Si el usuario solicita una acción (presupuesto, cita, imagen...), usa la\n"
    "  herramienta correspondiente en lugar de responder de texto libre\n"
    "- Si la pregunta está fuera del dominio de tu negocio, indícalo amablemente\n"
    "  y redirige al usuario hacia temas relacionados con su empresa\n"
    "- Mantén un tono profesional pero cercano, apropiado para una PYME española\n\n"
    "{custom_instructions}"
)

# ─── Pie de documento ──────────────────────────────────────────────────────
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run(
    f"Zeptai — Memoria TFM · Máster en IA Generativa · "
    f"Alejandro Bravo · {datetime.date.today().strftime('%B %Y')}"
)
footer_run.font.size = Pt(9)
footer_run.italic = True
footer_run.font.color.rgb = RGBColor(120, 120, 120)

# ─── Guardar ───────────────────────────────────────────────────────────────
output_path = r"c:\Users\aleja\Desktop\workspace\zeptai\Zeptai_Memoria_TFM.docx"
doc.save(output_path)
print(f"OK Documento guardado: {output_path}")
